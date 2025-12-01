"""
Routes pour la gestion des archives d'enquêtes
Permet d'archiver des enquêtes avec stockage de fichiers et de les consulter
"""
from flask import Blueprint, request, jsonify, send_file
import os
import datetime
import logging
from io import BytesIO
from models.models import Donnee
from models.models_enqueteur import DonneeEnqueteur
from models.enqueteur import Enqueteur
from models.enquete_archive import EnqueteArchive
from models.enquete_archive_file import EnqueteArchiveFile
from extensions import db

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

logger = logging.getLogger(__name__)

archives_bp = Blueprint('archives', __name__)

# Chemin de base pour le stockage des archives
BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
ARCHIVES_DIR = os.path.join(BASE_DIR, 'exports', 'archives')

def ensure_archives_directory():
    """S'assure que le dossier d'archives existe"""
    os.makedirs(ARCHIVES_DIR, exist_ok=True)

def generate_word_document(donnee):
    """
    Génère un document Word (.docx) pour une enquête
    """
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx n'est pas installé")
    
    doc = Document()
    
    # Récupérer les données enquêteur
    donnee_enqueteur = DonneeEnqueteur.query.filter_by(donnee_id=donnee.id).first()
    
    # Récupérer l'enquêteur
    enqueteur = None
    if donnee.enqueteurId:
        enqueteur = db.session.get(Enqueteur, donnee.enqueteurId)
    
    # Titre principal
    titre = doc.add_heading(f"Enquête n°{donnee.id} – {donnee.nom or ''} {donnee.prenom or ''}", level=1)
    titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in titre.runs:
        run.font.size = Pt(18)
        run.font.bold = True
    
    # Sous-titre
    date_str = donnee.updated_at.strftime('%d/%m/%Y') if donnee.updated_at else 'N/A'
    enqueteur_str = f"{enqueteur.nom} {enqueteur.prenom}" if enqueteur else "Non assigné"
    statut_str = donnee.statut_validation or "En attente"
    
    sous_titre = doc.add_paragraph(f"Date : {date_str} | Enquêteur : {enqueteur_str} | Statut : {statut_str}")
    sous_titre.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in sous_titre.runs:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(64, 64, 64)
    
    doc.add_paragraph()
    
    # Tableau des données
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    
    # En-tête
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Champ'
    hdr_cells[1].text = 'Valeur'
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(11)
    
    # Données de base
    fields_data = [
        ('N° Dossier', donnee.numeroDossier),
        ('Référence Dossier', donnee.referenceDossier),
        ('Type de Demande', donnee.typeDemande),
        ('Qualité', donnee.qualite),
        ('Nom', donnee.nom),
        ('Prénom', donnee.prenom),
        ('Date de Naissance', donnee.dateNaissance.strftime('%d/%m/%Y') if donnee.dateNaissance else None),
        ('Lieu de Naissance', donnee.lieuNaissance),
        ('Nom Patronymique', donnee.nomPatronymique),
        ('Adresse 1', donnee.adresse1),
        ('Adresse 2', donnee.adresse2),
        ('Ville', donnee.ville),
        ('Code Postal', donnee.codePostal),
        ('Pays de Résidence', donnee.paysResidence),
        ('Téléphone Personnel', donnee.telephonePersonnel),
        ('Nom Employeur', donnee.nomEmployeur),
    ]
    
    # Ajouter les données enquêteur
    if donnee_enqueteur:
        fields_data.extend([
            ('Code Résultat', donnee_enqueteur.code_resultat),
            ('Éléments Retrouvés', donnee_enqueteur.elements_retrouves),
            ('Date de Retour', donnee_enqueteur.date_retour.strftime('%d/%m/%Y') if donnee_enqueteur.date_retour else None),
        ])
    
    # Remplir le tableau
    for field_name, field_value in fields_data:
        if field_value:
            row_cells = table.add_row().cells
            row_cells[0].text = str(field_name)
            row_cells[1].text = str(field_value)
            
            for paragraph in row_cells[0].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)
            
            for paragraph in row_cells[1].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
    
    doc.add_paragraph()
    
    # Notes
    doc.add_heading('Notes / Commentaires', level=2)
    notes = donnee.commentaire or "Aucune note"
    if donnee_enqueteur and donnee_enqueteur.notes_personnelles:
        notes += f"\n\nNotes de l'enquêteur:\n{donnee_enqueteur.notes_personnelles}"
    
    notes_para = doc.add_paragraph(notes)
    for run in notes_para.runs:
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(64, 64, 64)
    
    return doc

@archives_bp.route('/api/archives/enquetes', methods=['GET'])
def get_archived_enquetes():
    """
    Récupère la liste des enquêtes archivées avec leurs métadonnées
    """
    try:
        page = request.args.get('page', 1, type=int)
        per_page = request.args.get('per_page', 50, type=int)
        
        # Récupérer les enquêtes archivées avec leurs fichiers
        archives_query = db.session.query(
            Donnee, DonneeEnqueteur, EnqueteArchiveFile
        ).join(
            DonneeEnqueteur, Donnee.id == DonneeEnqueteur.donnee_id, isouter=True
        ).join(
            EnqueteArchiveFile, Donnee.id == EnqueteArchiveFile.enquete_id, isouter=True
        ).filter(
            Donnee.statut_validation == 'archive'
        ).order_by(
            Donnee.updated_at.desc()
        )
        
        # Pagination
        archives_paginated = archives_query.paginate(page=page, per_page=per_page, error_out=False)
        
        result = []
        for donnee, donnee_enqueteur, archive_file in archives_paginated.items:
            # Récupérer l'enquêteur
            enqueteur = None
            if donnee.enqueteurId:
                enqueteur = db.session.get(Enqueteur, donnee.enqueteurId)
            
            enqueteur_nom = f"{enqueteur.nom} {enqueteur.prenom}" if enqueteur else "Non assigné"
            
            result.append({
                'id': donnee.id,
                'numeroDossier': donnee.numeroDossier,
                'referenceDossier': donnee.referenceDossier,
                'nom': donnee.nom,
                'prenom': donnee.prenom,
                'typeDemande': donnee.typeDemande,
                'enqueteurId': donnee.enqueteurId,
                'enqueteurNom': enqueteur_nom,
                'code_resultat': donnee_enqueteur.code_resultat if donnee_enqueteur else None,
                'elements_retrouves': donnee_enqueteur.elements_retrouves if donnee_enqueteur else None,
                'date_archive': donnee.updated_at.strftime('%Y-%m-%d %H:%M:%S') if donnee.updated_at else None,
                'archive_file_id': archive_file.id if archive_file else None,
                'filename': archive_file.filename if archive_file else None,
                'file_size': archive_file.file_size if archive_file else None,
                'utilisateur': archive_file.utilisateur if archive_file else None
            })
        
        return jsonify({
            'success': True,
            'data': result,
            'page': page,
            'per_page': per_page,
            'total': archives_paginated.total,
            'pages': archives_paginated.pages
        })
    except Exception as e:
        logger.error(f"Erreur lors de la récupération des archives: {str(e)}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@archives_bp.route('/api/archives/enquetes/<int:enquete_id>', methods=['GET'])
def get_archived_enquete_details(enquete_id):
    """
    Récupère les détails complets d'une enquête archivée (read-only)
    """
    try:
        donnee = db.session.get(Donnee, enquete_id)
        if not donnee:
            return jsonify({"error": "Enquête non trouvée"}), 404
        
        if donnee.statut_validation != 'archive':
            return jsonify({"error": "Cette enquête n'est pas archivée"}), 400
        
        # Récupérer les données enquêteur
        donnee_enqueteur = DonneeEnqueteur.query.filter_by(donnee_id=enquete_id).first()
        
        # Récupérer l'enquêteur
        enqueteur = None
        if donnee.enqueteurId:
            enqueteur = db.session.get(Enqueteur, donnee.enqueteurId)
        
        # Construire la réponse
        result = donnee.to_dict()
        
        if donnee_enqueteur:
            result['donnee_enqueteur'] = donnee_enqueteur.to_dict()
        
        if enqueteur:
            result['enqueteur'] = {
                'id': enqueteur.id,
                'nom': enqueteur.nom,
                'prenom': enqueteur.prenom,
                'email': enqueteur.email
            }
        
        return jsonify({
            'success': True,
            'data': result
        })
    except Exception as e:
        logger.error(f"Erreur lors de la récupération des détails de l'archive: {str(e)}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@archives_bp.route('/api/archives/enquetes/<int:enquete_id>/archive', methods=['POST'])
def archive_enquete(enquete_id):
    """
    Archive une enquête : génère le fichier Word, le stocke sur disque,
    met à jour le statut et crée les entrées dans les tables d'archives
    """
    try:
        donnee = db.session.get(Donnee, enquete_id)
        if not donnee:
            return jsonify({"error": "Enquête non trouvée"}), 404
        
        # Vérifier que l'enquête n'est pas déjà archivée
        if donnee.statut_validation == 'archive':
            return jsonify({"error": "Cette enquête est déjà archivée"}), 400
        
        # Vérifier qu'il y a des données enquêteur
        donnee_enqueteur = DonneeEnqueteur.query.filter_by(donnee_id=enquete_id).first()
        if not donnee_enqueteur or not donnee_enqueteur.code_resultat:
            return jsonify({"error": "L'enquête doit avoir un résultat avant d'être archivée"}), 400
        
        # S'assurer que le dossier d'archives existe
        ensure_archives_directory()
        
        # Générer le document Word
        doc = generate_word_document(donnee)
        
        # Créer un nom de fichier unique
        timestamp = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"Enquete_{donnee.numeroDossier}_{timestamp}.docx"
        
        # Créer le sous-dossier pour cette enquête
        enquete_dir = os.path.join(ARCHIVES_DIR, str(enquete_id))
        os.makedirs(enquete_dir, exist_ok=True)
        
        # Chemin complet du fichier
        filepath_full = os.path.join(enquete_dir, filename)
        
        # Sauvegarder le document sur disque
        doc.save(filepath_full)
        
        # Calculer la taille du fichier
        file_size = os.path.getsize(filepath_full)
        
        # Chemin relatif depuis le dossier exports/
        filepath_relative = os.path.join('archives', str(enquete_id), filename)
        
        # Récupérer l'utilisateur depuis la requête
        utilisateur = request.json.get('utilisateur', 'Administrateur') if request.json else 'Administrateur'
        
        # Mettre à jour le statut de l'enquête
        donnee.statut_validation = 'archive'
        donnee.updated_at = datetime.datetime.utcnow()
        
        # Créer l'entrée dans EnqueteArchive (pour compatibilité)
        archive = EnqueteArchive(
            enquete_id=enquete_id,
            nom_fichier=filename,
            utilisateur=utilisateur
        )
        db.session.add(archive)
        
        # Créer l'entrée dans EnqueteArchiveFile
        archive_file = EnqueteArchiveFile(
            enquete_id=enquete_id,
            filename=filename,
            filepath=filepath_relative,
            type_export='word',
            file_size=file_size,
            utilisateur=utilisateur
        )
        db.session.add(archive_file)
        
        db.session.commit()
        
        logger.info(f"Enquête {enquete_id} archivée avec succès. Fichier: {filepath_full}")
        
        return jsonify({
            'success': True,
            'message': 'Enquête archivée avec succès',
            'data': {
                'enquete_id': enquete_id,
                'archive_file_id': archive_file.id,
                'filename': filename,
                'file_size': file_size,
                'filepath': filepath_relative
            }
        })
        
    except Exception as e:
        db.session.rollback()
        logger.error(f"Erreur lors de l'archivage de l'enquête {enquete_id}: {str(e)}")
        return jsonify({
            'success': False,
            'error': f"Erreur lors de l'archivage: {str(e)}"
        }), 500

@archives_bp.route('/api/archives/enquetes/<int:archive_file_id>/download', methods=['GET'])
def download_archive_file(archive_file_id):
    """
    Télécharge un fichier archivé depuis le disque
    """
    try:
        archive_file = db.session.get(EnqueteArchiveFile, archive_file_id)
        if not archive_file:
            return jsonify({"error": "Archive non trouvée"}), 404
        
        # Construire le chemin complet du fichier
        filepath_full = os.path.join(BASE_DIR, 'exports', archive_file.filepath)
        
        # Vérifier que le fichier existe
        if not os.path.exists(filepath_full):
            logger.error(f"Fichier d'archive non trouvé: {filepath_full}")
            return jsonify({"error": "Fichier d'archive non trouvé sur le disque"}), 404
        
        # Envoyer le fichier
        return send_file(
            filepath_full,
            as_attachment=True,
            download_name=archive_file.filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        logger.error(f"Erreur lors du téléchargement de l'archive {archive_file_id}: {str(e)}")
        return jsonify({
            'success': False,
            'error': f"Erreur lors du téléchargement: {str(e)}"
        }), 500

def register_archives_routes(app):
    """Enregistre les routes d'archives dans l'application"""
    app.register_blueprint(archives_bp)
