# backend/routes/export.py

from flask import Blueprint, request, jsonify, send_file
import os
import datetime
import logging
import re
from io import StringIO, BytesIO
import tempfile
from models.models import Donnee, Fichier
from models.models_enqueteur import DonneeEnqueteur
from models.enqueteur import Enqueteur
from models.export_batch import ExportBatch
from extensions import db

# Configuration du logging
logger = logging.getLogger(__name__)

# Configuration pour l'export EOS
CODE_PRESTATAIRE = os.getenv('CODE_PRESTATAIRE', 'XXX')  # 3 lettres


# ========== FONCTIONS DE FORMATAGE POUR EXPORT EOS ==========

def clean_special_chars(text):
    """Nettoie les caractères spéciaux pour l'encodage cp1252"""
    if text is None:
        return ''
    
    # Remplacements pour les caractères problématiques
    replacements = {
        'œ': 'oe', 'Œ': 'OE',
        'æ': 'ae', 'Æ': 'AE',
        '€': 'EUR',
        '—': '-', '–': '-',
        ''': "'", ''': "'",
        '"': '"', '"': '"',
        '…': '...',
    }
    
    text = str(text)
    for old, new in replacements.items():
        text = text.replace(old, new)
    
    # Encoder en cp1252 avec remplacement des caractères non supportés
    try:
        text.encode('cp1252')
        return text
    except UnicodeEncodeError:
        # Remplacer les caractères non encodables par '?'
        return text.encode('cp1252', errors='replace').decode('cp1252')


def format_alphanum_eos(value, length):
    """Formate un champ alphanumérique : padding à droite avec des espaces"""
    if value is None:
        value = ''
    value_str = clean_special_chars(str(value).strip())
    if len(value_str) > length:
        value_str = value_str[:length]
    return value_str.ljust(length)


def format_numeric_eos(value, length):
    """
    Formate un champ numérique entier.
    - None → espaces (PAS zéros)
    - Valeurs négatives (-1) → justifié droite avec espaces
    - Valeurs positives → padding gauche avec zéros

    Exemples:
        format_numeric_eos(12, 3) → "012"
        format_numeric_eos(-1, 2) → "-1"
        format_numeric_eos(None, 3) → "   "
    """
    if value is None:
        return ' ' * length

    try:
        n = int(value)
        value_str = str(n)

        if len(value_str) > length:
            value_str = value_str[:length]

        # Si négatif, pas de zero-padding (garder le signe)
        if n < 0:
            return value_str.rjust(length)
        else:
            return value_str.rjust(length, '0')
    except (ValueError, TypeError):
        return ' ' * length


def format_date_eos(date_value):
    """Formate une date au format JJ/MM/AAAA (10 caractères)"""
    if date_value is None:
        return ' ' * 10
    
    if isinstance(date_value, str):
        try:
            date_value = datetime.datetime.strptime(date_value, '%Y-%m-%d').date()
        except:
            return ' ' * 10
    
    if isinstance(date_value, datetime.datetime):
        date_value = date_value.date()
    
    if isinstance(date_value, datetime.date):
        return date_value.strftime('%d/%m/%Y')
    
    return ' ' * 10


def format_montant_8(montant):
    """
    Formate un montant FACTURATION au format sur 8 chars.
    - Séparateur décimal: POINT (pas virgule)
    - Padding: ESPACES à gauche (pas zéros)
    - 2 décimales
    - Valeurs négatives acceptées

    Exemples (conforme ligne exemple):
        format_montant_8(24.00) → "   24.00"
        format_montant_8(0.00) → "    0.00"
        format_montant_8(None) → "    0.00"
        format_montant_8(-10.50) → "  -10.50"
    """
    if montant is None:
        montant = 0.0

    try:
        montant_float = float(montant)
        # Formater avec 2 décimales, POINT décimal
        montant_str = f"{montant_float:.2f}"

        # Right-justify avec espaces sur 8 chars
        return montant_str.rjust(8)
    except (ValueError, TypeError):
        return "    0.00"


def format_montant_10(montant):
    """
    Formate un montant REVENUS/SALAIRE au format sur 10 chars.
    - Séparateur décimal: POINT (cohérence avec montant_8)
    - Padding: ESPACES à gauche
    - 2 décimales
    - Valeurs négatives acceptées

    Exemples:
        format_montant_10(123456.78) → " 123456.78"
        format_montant_10(0.00) → "      0.00"
        format_montant_10(None) → "      0.00"
    """
    if montant is None:
        return ' ' * 10

    try:
        montant_float = float(montant)
        if montant_float == 0:
            return ' ' * 10

        # Formater avec 2 décimales, POINT décimal
        montant_str = f"{montant_float:.2f}"

        # Right-justify avec espaces sur 10 chars
        return montant_str.rjust(10)
    except (ValueError, TypeError):
        return ' ' * 10


def format_int_8(value):
    """
    Formate un entier sur 8 chars.
    - Padding: ESPACES à gauche (pas zéros)
    - Valeurs négatives acceptées
    - None → "       0"

    Exemples:
        format_int_8(0) → "       0"
        format_int_8(123) → "     123"
        format_int_8(-50) → "     -50"
        format_int_8(None) → "       0"
    """
    if value is None:
        value = 0

    try:
        n = int(value)
        value_str = str(n)
        return value_str.rjust(8)
    except (ValueError, TypeError):
        return "       0"


def get_tarif_from_eos(donnee, donnee_enqueteur, export_date):
    """
    Recherche le tarif actif dans tarifs_eos pour une enquête.

    Logique:
    1. Déterminer code_tarif:
       - Priorité 1: donnee.elementDemandes (trim)
       - Sinon: donnee_enqueteur.elements_retrouves (trim)
    2. Chercher montant actif à export_date dans tarifs_eos
    3. Si introuvable: retourner 0.00 + log warning

    Args:
        donnee: Objet Donnee
        donnee_enqueteur: Objet DonneeEnqueteur
        export_date: Date d'export (pour filtrer tarifs actifs)

    Returns:
        float: Montant tarif (0.00 si introuvable)
    """
    # Déterminer code_tarif
    code_tarif = None

    # Priorité 1: elementDemandes
    if hasattr(donnee, 'elementDemandes') and donnee.elementDemandes:
        code_tarif = str(donnee.elementDemandes).strip()

    # Sinon: elements_retrouves
    if not code_tarif and hasattr(donnee_enqueteur, 'elements_retrouves') and donnee_enqueteur.elements_retrouves:
        code_tarif = str(donnee_enqueteur.elements_retrouves).strip()

    if not code_tarif:
        logger.warning(f"Enquête ID={donnee.id}: aucun code tarif trouvé (elementDemandes et elements_retrouves vides)")
        return 0.0

    # Chercher tarif actif dans tarifs_eos
    try:
        from models.tarifs import TarifEOS

        tarif = TarifEOS.query.filter(
            TarifEOS.code == code_tarif,
            TarifEOS.actif == True,
            TarifEOS.date_debut <= export_date
        ).filter(
            db.or_(
                TarifEOS.date_fin.is_(None),
                TarifEOS.date_fin >= export_date
            )
        ).order_by(TarifEOS.date_debut.desc()).first()

        if tarif and tarif.montant is not None:
            logger.debug(f"Enquête ID={donnee.id}: tarif '{code_tarif}' trouvé → {tarif.montant}€")
            return float(tarif.montant)
        else:
            logger.warning(f"Enquête ID={donnee.id}: tarif '{code_tarif}' introuvable dans tarifs_eos pour date {export_date}")
            return 0.0

    except Exception as e:
        logger.error(f"Erreur recherche tarif pour enquête ID={donnee.id}: {e}")
        return 0.0


try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx n'est pas installé. L'export Word ne sera pas disponible.")

export_bp = Blueprint('export', __name__)

@export_bp.route('/api/export-enquetes', methods=['POST'])
def export_enquetes():
    """
    Génère un fichier Word (.docx) avec les résultats d'enquête
    - Page récapitulative au début (date réception + nombre de dossiers)
    - Une page par enquête avec toutes les données
    - Exporte UNIQUEMENT les enquêtes non encore exportées (exported=False)
    - Marque les enquêtes comme exportées après génération
    
    Utilisé depuis l'onglet "Données" pour exporter des enquêtes
    
    LIMITE : Maximum 1000 enquêtes par export pour éviter les surcharges
    """
    try:
        if not DOCX_AVAILABLE:
            return jsonify({
                'success': False,
                'error': 'python-docx n\'est pas installé'
            }), 500
        
        # MULTI-CLIENT: Récupérer le client (par défaut EOS)
        from client_utils import get_client_or_default
        data = request.get_json() or {}
        client_id_param = data.get('client_id')
        
        client = get_client_or_default(client_id=client_id_param)
        if not client:
            return jsonify({"success": False, "error": "Client introuvable"}), 404

        logger.info(f"Export Word demandé pour client: {client.code} (ID: {client.id})")

        # Limite maximale pour un seul export (performance et mémoire)
        MAX_EXPORT_LIMIT = 1000
        
        # Compter d'abord le nombre total d'enquêtes à exporter POUR CE CLIENT
        total_count = Donnee.query.filter(
            Donnee.client_id == client.id,
            Donnee.statut_validation.notin_(['validee', 'archivee']),
            Donnee.exported == False
        ).count()
        
        if total_count == 0:
            return jsonify({
                "success": False,
                "message": f"Aucune nouvelle enquête à exporter pour {client.nom}. Toutes les enquêtes ont déjà été exportées."
            }), 200
        
        # Avertir si dépassement de la limite
        if total_count > MAX_EXPORT_LIMIT:
            logger.warning(f"Export limité : {total_count} enquêtes disponibles, mais limite fixée à {MAX_EXPORT_LIMIT}")
        
        # Récupérer les enquêtes non exportées POUR CE CLIENT (AVEC LIMITE)
        donnees = Donnee.query.options(
            db.joinedload(Donnee.fichier)
        ).filter(
            Donnee.client_id == client.id,
            Donnee.statut_validation.notin_(['validee', 'archivee']),
            Donnee.exported == False
        ).order_by(Donnee.created_at.asc()).limit(MAX_EXPORT_LIMIT).all()
        
        if not donnees:
            return jsonify({
                "success": False,
                "message": "Aucune nouvelle enquête à exporter. Toutes les enquêtes ont déjà été exportées."
            }), 200
        
        logger.info(f"Génération d'un export Word pour {len(donnees)} enquête(s) sur {total_count} non exportée(s)")
        
        # Message si export partiel
        export_info = {
            'exported': len(donnees),
            'remaining': total_count - len(donnees),
            'total': total_count
        }
        
        # Extraire la date du NOM DU FICHIER (ex: LDMExp_20251120.txt -> 20/11/2025)
        date_reception = None
        if donnees and donnees[0].fichier:
            nom_fichier = donnees[0].fichier.nom
            # Format attendu: LDMExp_AAAAMMJJ.txt
            match = re.search(r'_(\d{8})', nom_fichier)
            if match:
                date_str = match.group(1)  # ex: 20251120
                try:
                    date_reception = datetime.datetime.strptime(date_str, '%Y%m%d')
                    logger.info(f"Date extraite du nom du fichier '{nom_fichier}': {date_reception.strftime('%d/%m/%Y')}")
                except:
                    logger.warning(f"Impossible de parser la date '{date_str}' du fichier '{nom_fichier}'")
                    date_reception = donnees[0].created_at
            else:
                logger.warning(f"Format de date non trouvé dans le nom du fichier '{nom_fichier}'")
                date_reception = donnees[0].created_at
        else:
            date_reception = donnees[0].created_at if donnees else datetime.datetime.now()
        
        nombre_dossiers = len(donnees)
        
        # Générer le document Word avec page récapitulative
        doc = generate_word_document_with_summary(donnees, date_reception, nombre_dossiers)
        
        # Marquer les enquêtes comme exportées
        now = datetime.datetime.now()
        for donnee in donnees:
            donnee.exported = True
            donnee.exported_at = now
        
        # Sauvegarder les modifications
        db.session.commit()
        logger.info(f"{len(donnees)} enquête(s) marquée(s) comme exportées")
        
        # Créer un fichier temporaire
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(temp_file.name)
        temp_file.close()
        
        # Envoyer le fichier au client
        return send_file(
            temp_file.name,
            as_attachment=True,
            download_name=f"Export_Nouvelles_Enquetes_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        db.session.rollback()
        logger.error(f"Erreur lors de l'export Word des enquêtes: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({"error": f"Erreur lors de l'export: {str(e)}"}), 500

def generate_export_content(donnees):
    """
    Génère le contenu du fichier d'export au format texte à longueur fixe
    selon les spécifications du cahier des charges EOS
    """
    output = StringIO()
    
    for donnee in donnees:
        # Récupérer les données enquêteur associées
        donnee_enqueteur = DonneeEnqueteur.query.filter_by(donnee_id=donnee.id).first()
        
        # Si pas de données enquêteur, on utilise des valeurs par défaut
        if not donnee_enqueteur:
            donnee_enqueteur = DonneeEnqueteur(
                donnee_id=donnee.id,
                client_id=donnee.client_id  # Hériter du client de la donnée
            )
        
        # Récupérer l'enquêteur
        enqueteur = None
        if donnee.enqueteurId:
            enqueteur = db.session.get(Enqueteur, donnee.enqueteurId)
        
        # Construire la ligne selon le format du cahier des charges
        line = format_export_line(donnee, donnee_enqueteur, enqueteur)
        output.write(line + '\n')
    
    return output.getvalue()

def format_export_line(donnee, donnee_enqueteur, enqueteur=None):
    """
    Formatte une ligne de données selon le format spécifié dans le cahier des charges
    Prend en compte les corrections d'état civil lorsque le flag est activé
    """
    # Initialisation de la ligne avec des espaces
    line = " " * 1854  # Longueur totale de la ligne selon le cahier des charges
    
    # Liste des positions et longueurs de chaque champ selon le cahier des charges
    fields = [
        # Champ, début, longueur
        ("numeroDossier", 0, 10),
        ("referenceDossier", 10, 15),
        ("numeroInterlocuteur", 25, 12),
        ("guidInterlocuteur", 37, 36),
        ("typeDemande", 73, 3),
        ("numeroDemande", 76, 11),
        ("numeroDemandeContestee", 87, 11),
        ("numeroDemandeInitiale", 98, 11),
        ("forfaitDemande", 109, 16),
        ("dateRetourEspere", 125, 10),
        ("qualite", 135, 10),
        ("nom", 145, 30),
        ("prenom", 175, 20),
        ("dateNaissance", 195, 10),
        ("lieuNaissance", 205, 50),
        ("codePostalNaissance", 255, 10),
        ("paysNaissance", 265, 32),
        ("nomPatronymique", 297, 30),
        ("dateRetour", 327, 10),
        ("codeResultat", 337, 1),
        ("elementsRetrouves", 338, 10),
        ("flagEtatCivilErrone", 348, 1),
        ("numeroFacture", 349, 9),
        ("dateFacture", 358, 10),
        ("montantFacture", 368, 8),
        ("tarifApplique", 376, 8),
        ("cumulMontantsPrecedents", 384, 8),
        ("repriseFacturation", 392, 8),
        ("remiseEventuelle", 400, 8),
        ("dateDeces", 408, 10),
        ("numeroActeDeces", 418, 10),
        ("codeInseeDeces", 428, 5),
        ("codePostalDeces", 433, 10),
        ("localiteDeces", 443, 32),
        ("adresse1", 475, 32),
        ("adresse2", 507, 32),
        ("adresse3", 539, 32),
        ("adresse4", 571, 32),
        ("codePostal", 603, 10),
        ("ville", 613, 32),
        ("paysResidence", 645, 32),
        ("telephonePersonnel", 677, 15),
        ("telephoneEmployeur", 692, 15),
        ("nomEmployeur", 707, 32),
        ("telephoneEmployeur2", 739, 15),
        ("telecopieEmployeur", 754, 15),
        ("adresse1Employeur", 769, 32),
        ("adresse2Employeur", 801, 32),
        ("adresse3Employeur", 833, 32),
        ("adresse4Employeur", 865, 32),
        ("codePostalEmployeur", 897, 10),
        ("villeEmployeur", 907, 32),
        ("paysEmployeur", 939, 32),
        ("banqueDomiciliation", 971, 32),
        ("libelleGuichet", 1003, 30),
        ("titulaireCompte", 1033, 32),
        ("codeBanque", 1065, 5),
        ("codeGuichet", 1070, 5),
        ("numeroCompte", 1075, 11),
        ("ribCompte", 1086, 2),
    ]
    
    # Déterminer si on doit utiliser les données d'état civil corrigées
    use_corrected_data = donnee_enqueteur and hasattr(donnee_enqueteur, 'flag_etat_civil_errone') and donnee_enqueteur.flag_etat_civil_errone == 'E'
    
    # Déterminer les valeurs d'état civil à utiliser (originales ou corrigées)
    qualite = donnee.qualite
    nom = donnee.nom
    prenom = donnee.prenom
    nom_patronymique = donnee.nomPatronymique
    date_naissance = donnee.dateNaissance.strftime('%d/%m/%Y') if donnee.dateNaissance else ""
    lieu_naissance = donnee.lieuNaissance
    code_postal_naissance = donnee.codePostalNaissance
    pays_naissance = donnee.paysNaissance
    
    # Si on doit utiliser les données corrigées, on les récupère
    if use_corrected_data:
        if hasattr(donnee_enqueteur, 'qualite_corrigee') and donnee_enqueteur.qualite_corrigee:
            qualite = donnee_enqueteur.qualite_corrigee
        if hasattr(donnee_enqueteur, 'nom_corrige') and donnee_enqueteur.nom_corrige:
            nom = donnee_enqueteur.nom_corrige
        if hasattr(donnee_enqueteur, 'prenom_corrige') and donnee_enqueteur.prenom_corrige:
            prenom = donnee_enqueteur.prenom_corrige
        if hasattr(donnee_enqueteur, 'nom_patronymique_corrige') and donnee_enqueteur.nom_patronymique_corrige:
            nom_patronymique = donnee_enqueteur.nom_patronymique_corrige
        if hasattr(donnee_enqueteur, 'date_naissance_corrigee') and donnee_enqueteur.date_naissance_corrigee:
            date_naissance = donnee_enqueteur.date_naissance_corrigee.strftime('%d/%m/%Y')
        if hasattr(donnee_enqueteur, 'lieu_naissance_corrige') and donnee_enqueteur.lieu_naissance_corrige:
            lieu_naissance = donnee_enqueteur.lieu_naissance_corrige
        if hasattr(donnee_enqueteur, 'code_postal_naissance_corrige') and donnee_enqueteur.code_postal_naissance_corrige:
            code_postal_naissance = donnee_enqueteur.code_postal_naissance_corrige
        if hasattr(donnee_enqueteur, 'pays_naissance_corrige') and donnee_enqueteur.pays_naissance_corrige:
            pays_naissance = donnee_enqueteur.pays_naissance_corrige
    
    # Formater la date de retour avec la date actuelle
    date_retour = datetime.datetime.now().strftime('%d/%m/%Y')
    
    # Préparer toutes les valeurs à insérer
    values = {
        "numeroDossier": donnee.numeroDossier or "",
        "referenceDossier": donnee.referenceDossier or "",
        "numeroInterlocuteur": donnee.numeroInterlocuteur or "",
        "guidInterlocuteur": donnee.guidInterlocuteur or "",
        "typeDemande": donnee.typeDemande or "",
        "numeroDemande": donnee.numeroDemande or "",
        "numeroDemandeContestee": donnee.numeroDemandeContestee or "",
        "numeroDemandeInitiale": donnee.numeroDemandeInitiale or "",
        "forfaitDemande": donnee.forfaitDemande or "",
        "dateRetourEspere": donnee.dateRetourEspere.strftime('%d/%m/%Y') if donnee.dateRetourEspere else "",
        
        # État civil (original ou corrigé)
        "qualite": qualite or "",
        "nom": nom or "",
        "prenom": prenom or "",
        "dateNaissance": date_naissance or "",
        "lieuNaissance": lieu_naissance or "",
        "codePostalNaissance": code_postal_naissance or "",
        "paysNaissance": pays_naissance or "",
        "nomPatronymique": nom_patronymique or "",
        
        "dateRetour": date_retour,
        "codeResultat": donnee_enqueteur.code_resultat or "",
        "elementsRetrouves": donnee_enqueteur.elements_retrouves or "",
        "flagEtatCivilErrone": donnee_enqueteur.flag_etat_civil_errone if hasattr(donnee_enqueteur, 'flag_etat_civil_errone') else "",
        "numeroFacture": "",  # Laisser vide selon le cahier des charges
        "dateFacture": "",    # Laisser vide selon le cahier des charges
        "montantFacture": "0",
        "tarifApplique": "0",
        "cumulMontantsPrecedents": "0",
        "repriseFacturation": "0",
        "remiseEventuelle": "0",
        "dateDeces": donnee_enqueteur.date_deces.strftime('%d/%m/%Y') if hasattr(donnee_enqueteur, 'date_deces') and donnee_enqueteur.date_deces else "",
        "numeroActeDeces": donnee_enqueteur.numero_acte_deces or "",
        "codeInseeDeces": donnee_enqueteur.code_insee_deces or "",
        "codePostalDeces": donnee_enqueteur.code_postal_deces or "",
        "localiteDeces": donnee_enqueteur.localite_deces or "",
        "adresse1": donnee_enqueteur.adresse1 or "",
        "adresse2": donnee_enqueteur.adresse2 or "",
        "adresse3": donnee_enqueteur.adresse3 or "",
        "adresse4": donnee_enqueteur.adresse4 or "",
        "codePostal": donnee_enqueteur.code_postal or "",
        "ville": donnee_enqueteur.ville or "",
        "paysResidence": donnee_enqueteur.pays_residence or "",
        "telephonePersonnel": donnee_enqueteur.telephone_personnel or "",
        "telephoneEmployeur": donnee_enqueteur.telephone_chez_employeur or "",
        "nomEmployeur": donnee_enqueteur.nom_employeur or "",
        "telephoneEmployeur2": donnee_enqueteur.telephone_employeur or "",
        "telecopieEmployeur": donnee_enqueteur.telecopie_employeur or "",
        "adresse1Employeur": donnee_enqueteur.adresse1_employeur or "",
        "adresse2Employeur": donnee_enqueteur.adresse2_employeur or "",
        "adresse3Employeur": donnee_enqueteur.adresse3_employeur or "",
        "adresse4Employeur": donnee_enqueteur.adresse4_employeur or "",
        "codePostalEmployeur": donnee_enqueteur.code_postal_employeur or "",
        "villeEmployeur": donnee_enqueteur.ville_employeur or "",
        "paysEmployeur": donnee_enqueteur.pays_employeur or "",
        "banqueDomiciliation": donnee_enqueteur.banque_domiciliation or "",
        "libelleGuichet": donnee_enqueteur.libelle_guichet or "",
        "titulaireCompte": donnee_enqueteur.titulaire_compte or "",
        "codeBanque": donnee_enqueteur.code_banque or "",
        "codeGuichet": donnee_enqueteur.code_guichet or "",
        "numeroCompte": "",  # Laisser vide selon le cahier des charges
        "ribCompte": "",     # Laisser vide selon le cahier des charges
    }
    
    # Construire la ligne en insérant chaque valeur à la position correcte
    line_list = list(line)
    for field, start, length in fields:
        value = values.get(field, "")
        # S'assurer que value est une chaîne de caractères
        if value is None:
            value = ""
        else:
            value = str(value)
        # Tronquer si la valeur est trop longue
        if len(value) > length:
            value = value[:length]
        # Insérer la valeur dans la ligne à la position correcte
        for i, char in enumerate(value):
            if i < length and start + i < len(line_list):
                line_list[start + i] = char
    
    # Reconvertir la liste en chaîne
    return "".join(line_list)


def generate_word_document_with_summary(donnees, date_reception, nombre_dossiers):
    """
    Génère un document Word : UNE PAGE PAR ENQUÊTE
    Date de réception et nombre de dossiers affichés sur chaque page
    
    Args:
        donnees: Liste des enquêtes à exporter
        date_reception: Date de réception des dossiers  
        nombre_dossiers: Nombre total de dossiers exportés
    """
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx n'est pas installé")
    
    doc = Document()
    
    # ========== UNE PAGE PAR ENQUÊTE ==========
    for i, donnee in enumerate(donnees):
        if i > 0:
            doc.add_page_break()
        
        generate_enquete_page(doc, donnee, i + 1, len(donnees), date_reception, nombre_dossiers)
    
    return doc


def generate_enquete_page(doc, donnee, numero_enquete, total_enquetes, date_reception, nombre_dossiers):
    """
    Génère une page pour une seule enquête dans le document Word
    UNE ENQUÊTE = UNE PAGE avec date réception et nombre de dossiers en haut
    
    Args:
        doc: Document Word
        donnee: Enquête à afficher
        numero_enquete: Numéro de l'enquête dans le lot (pour affichage)
        total_enquetes: Total d'enquêtes dans le lot
        date_reception: Date de réception des dossiers
        nombre_dossiers: Nombre total de dossiers exportés
    """
    # Récupérer les données enquêteur
    donnee_enqueteur = DonneeEnqueteur.query.filter_by(donnee_id=donnee.id).first()
    
    # Titre principal - COMPACT
    titre = doc.add_heading(f"ENQUÊTE {numero_enquete}/{total_enquetes} - N°{donnee.id}", level=1)
    titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titre.paragraph_format.space_after = Pt(4)  # Espacement réduit
    for run in titre.runs:
        run.font.size = Pt(16)  # Plus petit
        run.font.bold = True
    
    # Informations récapitulatives EN HAUT DE CHAQUE PAGE - COMPACT
    recap_para = doc.add_paragraph()
    recap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    recap_para.paragraph_format.space_after = Pt(4)  # Espacement réduit
    run1 = recap_para.add_run(f"Date: {date_reception.strftime('%d/%m/%Y') if date_reception else 'N/A'}")
    run1.font.size = Pt(16)  # Plus petit
    run1.font.bold = True
    recap_para.add_run(" | ")
    run2 = recap_para.add_run(f"Dossiers: {nombre_dossiers}")
    run2.font.size = Pt(16)  # Plus petit
    run2.font.bold = True
    run2.font.color.rgb = RGBColor(0, 102, 204)
    
    # Fonction helper pour ajouter des sections (FORMAT COMPACT POUR 1 PAGE)
    def add_table_section(doc, title, fields):
        """Ajoute une section avec un tableau de données - Format compact"""
        doc.add_heading(title, level=2)
        for run in doc.paragraphs[-1].runs:
            run.font.size = Pt(16)  # Titre plus petit
        
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Light Grid Accent 1'
        
        # En-tête
        hdr = table.rows[0].cells
        hdr[0].text = 'Champ'
        hdr[1].text = 'Valeur'
        for cell in hdr:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(16)  # Police plus petite
        
        # Données
        has_data = False
        for field_name, field_value in fields:
            if field_value not in [None, '', []]:
                has_data = True
                row = table.add_row().cells
                row[0].text = str(field_name)
                row[1].text = str(field_value)
                # Style
                for paragraph in row[0].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(16)
                for paragraph in row[1].paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(16)
        
        if not has_data:
            row = table.add_row().cells
            row[0].text = ''
            row[1].text = 'Aucune donnée'
        
        doc.add_paragraph()  # Espacement
        return has_data
    
    # ========== FORMAT COMPACT : UNE ENQUÊTE = UNE PAGE ==========
    
    # Tableau unique avec les informations essentielles
    doc.add_heading('Informations de l\'enquête', level=2)
    
    # Style compact avec taille de police réduite
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    
    # En-tête
    hdr = table.rows[0].cells
    hdr[0].text = 'Champ'
    hdr[1].text = 'Valeur'
    for cell in hdr:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(16)
    
    # Fonction pour ajouter une ligne
    def add_row(label, value):
        if value not in [None, '', 'None']:
            row = table.add_row().cells
            row[0].text = str(label)
            row[1].text = str(value)
            for paragraph in row[0].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(16)
            for paragraph in row[1].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(16)
    
    # TOUTES LES DONNÉES DU FICHIER (Format compact pour 1 page)
    
    # 1. Identification complète
    add_row('N° Dossier', donnee.numeroDossier)
    add_row('Type Demande', donnee.typeDemande)
    add_row('N° Demande', donnee.numeroDemande)
    add_row('N° Demande Contestée', donnee.numeroDemandeContestee)
    add_row('N° Demande Initiale', donnee.numeroDemandeInitiale)
    add_row('Forfait', donnee.forfaitDemande)
    add_row('Date Envoi', donnee.datedenvoie.strftime('%d/%m/%Y') if donnee.datedenvoie else None)
    add_row('Date Retour Espéré', donnee.dateRetourEspere.strftime('%d/%m/%Y') if donnee.dateRetourEspere else None)
    add_row('Date Butoir', donnee.date_butoir.strftime('%d/%m/%Y') if donnee.date_butoir else None)
    add_row('Code Société', donnee.codesociete)
    add_row('Urgence', donnee.urgence)
    
    # 2. État Civil complet
    add_row('Qualité', donnee.qualite)
    add_row('Nom', donnee.nom)
    add_row('Prénom', donnee.prenom)
    add_row('Nom Patronymique', donnee.nomPatronymique)
    add_row('Date Naissance', donnee.dateNaissance.strftime('%d/%m/%Y') if donnee.dateNaissance else None)
    add_row('Lieu Naissance', donnee.lieuNaissance)
    add_row('CP Naissance', donnee.codePostalNaissance)
    add_row('Pays Naissance', donnee.paysNaissance)
    
    # 3. Adresse complète
    add_row('Adresse 1', donnee.adresse1)
    add_row('Adresse 2', donnee.adresse2)
    add_row('Adresse 3', donnee.adresse3)
    add_row('Adresse 4', donnee.adresse4)
    add_row('Code Postal', donnee.codePostal)
    add_row('Ville', donnee.ville)
    add_row('Pays', donnee.paysResidence)
    add_row('Tél Personnel', donnee.telephonePersonnel)
    
    # 4. Employeur initial complet
    add_row('Employeur', donnee.nomEmployeur)
    add_row('Tél Employeur', donnee.telephoneEmployeur)
    add_row('Fax Employeur', donnee.telecopieEmployeur)
    
    # 5. Banque initiale complète
    add_row('Banque', donnee.banqueDomiciliation)
    add_row('Libellé Guichet', donnee.libelleGuichet)

    
    # 6. Éléments demandés
    if donnee.client and donnee.client.code == 'PARTNER':
        add_row('Éléments Demandés ', donnee.recherche)
        add_row('Instructions ', donnee.instructions)
    else:
        add_row('Éléments Demandés', donnee.elementDemandes)
    
    add_row('Éléments Obligatoires', donnee.elementObligatoires)
    add_row('Éléments Contestés', donnee.elementContestes)
    add_row('Code Motif', donnee.codeMotif)
    add_row('Motif Contestation', donnee.motifDeContestation)
    
    # 7. Commentaire complet
    if donnee.commentaire:
        add_row('Commentaire', donnee.commentaire)
    
    # === RÉSULTATS ENQUÊTEUR (TOUS) ===
    if donnee_enqueteur:
        add_row('━━━━━━━━━━', '━━━━━━━━━━')  # Séparateur
        add_row('CODE RÉSULTAT', donnee_enqueteur.code_resultat)
        add_row('Éléments Retrouvés', donnee_enqueteur.elements_retrouves)
        add_row('Date Retour', donnee_enqueteur.date_retour.strftime('%d/%m/%Y') if donnee_enqueteur.date_retour else None)
        add_row('État Civil Erroné', donnee_enqueteur.flag_etat_civil_errone)
        
        # Adresse trouvée complète
        add_row('Adr1 Trouvée', donnee_enqueteur.adresse1)
        add_row('Adr2 Trouvée', donnee_enqueteur.adresse2)
        add_row('Adr3 Trouvée', donnee_enqueteur.adresse3)
        add_row('Adr4 Trouvée', donnee_enqueteur.adresse4)
        add_row('CP Trouvé', donnee_enqueteur.code_postal)
        add_row('Ville Trouvée', donnee_enqueteur.ville)
        add_row('Pays Trouvé', donnee_enqueteur.pays_residence)
        add_row('Tél Trouvé', donnee_enqueteur.telephone_personnel)
        add_row('Tél Employeur Trouvé', donnee_enqueteur.telephone_chez_employeur)
        
        # Décès
        add_row('Date Décès', donnee_enqueteur.date_deces.strftime('%d/%m/%Y') if donnee_enqueteur.date_deces else None)
        add_row('N° Acte Décès', donnee_enqueteur.numero_acte_deces)
        add_row('Code INSEE Décès', donnee_enqueteur.code_insee_deces)
        add_row('Localité Décès', donnee_enqueteur.localite_deces)
        
        # Employeur trouvé complet
        add_row('Employeur Trouvé', donnee_enqueteur.nom_employeur)
        add_row('Tél Employeur T.', donnee_enqueteur.telephone_employeur)
        add_row('Fax Employeur T.', donnee_enqueteur.telecopie_employeur)
        add_row('Adr1 Employeur', donnee_enqueteur.adresse1_employeur)
        add_row('Adr2 Employeur', donnee_enqueteur.adresse2_employeur)
        add_row('Adr3 Employeur', donnee_enqueteur.adresse3_employeur)
        add_row('Adr4 Employeur', donnee_enqueteur.adresse4_employeur)
        add_row('CP Employeur', donnee_enqueteur.code_postal_employeur)
        add_row('Ville Employeur', donnee_enqueteur.ville_employeur)
        
        # Banque trouvée complète
        add_row('Banque Trouvée', donnee_enqueteur.banque_domiciliation)
        add_row('Libellé Guichet T.', donnee_enqueteur.libelle_guichet)
        add_row('Titulaire Compte T.', donnee_enqueteur.titulaire_compte)
        add_row('Code Banque T.', donnee_enqueteur.code_banque)
        add_row('Code Guichet T.', donnee_enqueteur.code_guichet)
        
        # Revenus (tous les 3)
        add_row('Montant Revenu 1', f"{donnee_enqueteur.montant_revenu1:.2f} €" if donnee_enqueteur.montant_revenu1 else None)
        add_row('Période Revenu 1', donnee_enqueteur.periode_versement_revenu1)
        add_row('Fréquence Revenu 1', donnee_enqueteur.frequence_versement_revenu1)
        add_row('Montant Revenu 2', f"{donnee_enqueteur.montant_revenu2:.2f} €" if donnee_enqueteur.montant_revenu2 else None)
        add_row('Période Revenu 2', donnee_enqueteur.periode_versement_revenu2)
        add_row('Fréquence Revenu 2', donnee_enqueteur.frequence_versement_revenu2)
        add_row('Montant Revenu 3', f"{donnee_enqueteur.montant_revenu3:.2f} €" if donnee_enqueteur.montant_revenu3 else None)
        add_row('Période Revenu 3', donnee_enqueteur.periode_versement_revenu3)
        add_row('Fréquence Revenu 3', donnee_enqueteur.frequence_versement_revenu3)
        
        # Mémos (TOUS les 5)
        for i in range(1, 6):
            memo_attr = f'memo{i}'
            if hasattr(donnee_enqueteur, memo_attr):
                memo_value = getattr(donnee_enqueteur, memo_attr)
                if memo_value:
                    add_row(f'Mémo {i}', memo_value)
        
        # Notes personnelles
        if hasattr(donnee_enqueteur, 'notes_personnelles') and donnee_enqueteur.notes_personnelles:
            add_row('Notes', donnee_enqueteur.notes_personnelles)


def generate_word_document(donnees):
    """
    Génère un document Word (.docx) contenant les enquêtes.
    Chaque enquête est sur une nouvelle page avec TOUTES les données de la table.
    """
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx n'est pas installé")
    
    doc = Document()
    
    def add_table_section(doc, title, fields):
        """Ajoute une section avec un tableau de données"""
        doc.add_heading(title, level=2)
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Light Grid Accent 1'
        
        # En-tête
        hdr = table.rows[0].cells
        hdr[0].text = 'Champ'
        hdr[1].text = 'Valeur'
        for cell in hdr:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(16)
        
        # Données
        has_data = False
        for field_name, field_value in fields:
            if field_value not in [None, '', []]:
                has_data = True
                row = table.add_row().cells
                row[0].text = str(field_name)
                row[1].text = str(field_value)
                # Style
                for paragraph in row[0].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(16)
                for paragraph in row[1].paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(16)
        
        if not has_data:
            row = table.add_row().cells
            row[0].text = ''
            row[1].text = 'Aucune donnée'
        
        doc.add_paragraph()  # Espacement
        return has_data
    
    for i, donnee in enumerate(donnees):
        # Ajouter un saut de page entre les enquêtes (sauf pour la première)
        if i > 0:
            doc.add_page_break()
        
        # Récupérer les données enquêteur
        donnee_enqueteur = DonneeEnqueteur.query.filter_by(donnee_id=donnee.id).first()
        
        # Récupérer l'enquêteur
        enqueteur = None
        if donnee.enqueteurId:
            enqueteur = db.session.get(Enqueteur, donnee.enqueteurId)
        
        # Titre principal
        titre = doc.add_heading(f"ENQUÊTE N° {donnee.id}", level=1)
        titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in titre.runs:
            run.font.size = Pt(18)
            run.font.bold = True
        
        # Sous-titre
        date_str = donnee.updated_at.strftime('%d/%m/%Y %H:%M') if donnee.updated_at else 'N/A'
        enqueteur_str = f"{enqueteur.nom} {enqueteur.prenom}" if enqueteur else "Non assigné"
        statut_str = donnee.statut_validation or "En attente"
        
        sous_titre = doc.add_paragraph(f"Dernière mise à jour : {date_str} | Enquêteur : {enqueteur_str} | Statut : {statut_str}")
        sous_titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in sous_titre.runs:
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(64, 64, 64)
        
        doc.add_paragraph()
        
        # ========== SECTION 1: IDENTIFICATION DU DOSSIER ==========
        add_table_section(doc, '1. Identification du Dossier', [
            ('N° Dossier', donnee.numeroDossier),
            ('Référence Dossier', donnee.referenceDossier),
            ('N° Interlocuteur', donnee.numeroInterlocuteur),
            ('GUID Interlocuteur', donnee.guidInterlocuteur),
            ('Type de Demande', donnee.typeDemande),
            ('N° Demande', donnee.numeroDemande),
            ('N° Demande Contestée', donnee.numeroDemandeContestee),
            ('N° Demande Initiale', donnee.numeroDemandeInitiale),
            ('Forfait Demande', donnee.forfaitDemande),
            ('Date d\'Envoi', donnee.datedenvoie.strftime('%d/%m/%Y') if donnee.datedenvoie else None),
            ('Date Retour Espéré', donnee.dateRetourEspere.strftime('%d/%m/%Y') if donnee.dateRetourEspere else None),
            ('Date Butoir', donnee.date_butoir.strftime('%d/%m/%Y') if donnee.date_butoir else None),
            ('Code Société', donnee.codesociete),
            ('Urgence', donnee.urgence),
        ])
        
        # ========== SECTION 2: ÉTAT CIVIL ==========
        add_table_section(doc, '2. État Civil', [
            ('Qualité', donnee.qualite),
            ('Nom', donnee.nom),
            ('Prénom', donnee.prenom),
            ('Nom Patronymique', donnee.nomPatronymique),
            ('Date de Naissance', donnee.dateNaissance.strftime('%d/%m/%Y') if donnee.dateNaissance else None),
            ('Lieu de Naissance', donnee.lieuNaissance),
            ('Code Postal Naissance', donnee.codePostalNaissance),
            ('Pays de Naissance', donnee.paysNaissance),
        ])
        
        # ========== SECTION 3: ADRESSE PERSONNELLE ==========
        add_table_section(doc, '3. Adresse Personnelle', [
            ('Adresse 1', donnee.adresse1),
            ('Adresse 2', donnee.adresse2),
            ('Adresse 3', donnee.adresse3),
            ('Adresse 4', donnee.adresse4),
            ('Code Postal', donnee.codePostal),
            ('Ville', donnee.ville),
            ('Pays de Résidence', donnee.paysResidence),
            ('Téléphone Personnel', donnee.telephonePersonnel),
        ])
        
        # ========== SECTION 4: INFORMATIONS EMPLOYEUR (DONNÉES INITIALES) ==========
        add_table_section(doc, '4. Informations Employeur (Données Initiales)', [
            ('Nom Employeur', donnee.nomEmployeur),
            ('Téléphone Employeur', donnee.telephoneEmployeur),
            ('Télécopie Employeur', donnee.telecopieEmployeur),
        ])
        
        # ========== SECTION 5: INFORMATIONS BANCAIRES (DONNÉES INITIALES) ==========
        add_table_section(doc, '5. Informations Bancaires (Données Initiales)', [
            ('Banque de Domiciliation', donnee.banqueDomiciliation),
            ('Libellé Guichet', donnee.libelleGuichet),
            ('Titulaire du Compte', donnee.titulaireCompte),
            ('Code Banque', donnee.codeBanque),
            ('Code Guichet', donnee.codeGuichet),
            ('N° Compte', donnee.numeroCompte),
            ('Clé RIB', donnee.ribCompte),
        ])
        
        # ========== SECTION 6: ÉLÉMENTS DEMANDÉS ET CONTESTATION ==========
        add_table_section(doc, '6. Éléments Demandés et Contestation', [
            ('Éléments Demandés', donnee.elementDemandes),
            ('Éléments Obligatoires', donnee.elementObligatoires),
            ('Éléments Contestés', donnee.elementContestes),
            ('Code Motif', donnee.codeMotif),
            ('Motif de Contestation', donnee.motifDeContestation),
            ('Est une Contestation', 'Oui' if donnee.est_contestation else 'Non'),
            ('Date Contestation', donnee.date_contestation.strftime('%d/%m/%Y') if donnee.date_contestation else None),
            ('Code Motif Contestation', donnee.motif_contestation_code),
            ('Détail Motif Contestation', donnee.motif_contestation_detail),
        ])
        
        # ========== SECTION 7: INFORMATIONS FINANCIÈRES ==========
        add_table_section(doc, '7. Informations Financières', [
            ('Cumul Montants Précédents', f"{donnee.cumulMontantsPrecedents:.2f} €" if donnee.cumulMontantsPrecedents else None),
        ])
        
        # ========== SECTION 8: COMMENTAIRE INITIAL ==========
        if donnee.commentaire:
            doc.add_heading('8. Commentaire Initial', level=2)
            para = doc.add_paragraph(donnee.commentaire)
            for run in para.runs:
                run.font.size = Pt(16)
            doc.add_paragraph()
        
        # ========== SECTIONS ENQUÊTEUR (SI DONNÉES DISPONIBLES) ==========
        if donnee_enqueteur:
            # Section 9: Résultat de l'enquête
            add_table_section(doc, '9. Résultat de l\'Enquête', [
                ('Code Résultat', donnee_enqueteur.code_resultat),
                ('Éléments Retrouvés', donnee_enqueteur.elements_retrouves),
                ('Date de Retour', donnee_enqueteur.date_retour.strftime('%d/%m/%Y') if donnee_enqueteur.date_retour else None),
                ('Flag État Civil Erroné', donnee_enqueteur.flag_etat_civil_errone),
            ])
            
            # Section 10: État civil corrigé
            if donnee_enqueteur.flag_etat_civil_errone:
                add_table_section(doc, '10. État Civil Corrigé', [
                    ('Qualité Corrigée', donnee_enqueteur.qualite_corrigee),
                    ('Nom Corrigé', donnee_enqueteur.nom_corrige),
                    ('Prénom Corrigé', donnee_enqueteur.prenom_corrige),
                    ('Nom Patronymique Corrigé', donnee_enqueteur.nom_patronymique_corrige),
                    ('Date Naissance Corrigée', donnee_enqueteur.date_naissance_corrigee.strftime('%d/%m/%Y') if donnee_enqueteur.date_naissance_corrigee else None),
                    ('Lieu Naissance Corrigé', donnee_enqueteur.lieu_naissance_corrige),
                    ('CP Naissance Corrigé', donnee_enqueteur.code_postal_naissance_corrige),
                    ('Pays Naissance Corrigé', donnee_enqueteur.pays_naissance_corrige),
                    ('Type de Divergence', donnee_enqueteur.type_divergence),
                ])
            
            # Section 11: Adresse trouvée par l'enquêteur
            add_table_section(doc, '11. Adresse Trouvée', [
                ('Adresse 1', donnee_enqueteur.adresse1),
                ('Adresse 2', donnee_enqueteur.adresse2),
                ('Adresse 3', donnee_enqueteur.adresse3),
                ('Adresse 4', donnee_enqueteur.adresse4),
                ('Code Postal', donnee_enqueteur.code_postal),
                ('Ville', donnee_enqueteur.ville),
                ('Pays', donnee_enqueteur.pays_residence),
                ('Téléphone Personnel', donnee_enqueteur.telephone_personnel),
                ('Téléphone chez Employeur', donnee_enqueteur.telephone_chez_employeur),
            ])
            
            # Section 12: Employeur (données enquêteur)
            add_table_section(doc, '12. Informations Employeur Trouvées', [
                ('Nom Employeur', donnee_enqueteur.nom_employeur),
                ('Téléphone', donnee_enqueteur.telephone_employeur),
                ('Télécopie', donnee_enqueteur.telecopie_employeur),
                ('Adresse 1', donnee_enqueteur.adresse1_employeur),
                ('Adresse 2', donnee_enqueteur.adresse2_employeur),
                ('Adresse 3', donnee_enqueteur.adresse3_employeur),
                ('Adresse 4', donnee_enqueteur.adresse4_employeur),
                ('Code Postal', donnee_enqueteur.code_postal_employeur),
                ('Ville', donnee_enqueteur.ville_employeur),
                ('Pays', donnee_enqueteur.pays_employeur),
            ])
            
            # Section 13: Informations bancaires trouvées
            add_table_section(doc, '13. Informations Bancaires Trouvées', [
                ('Banque de Domiciliation', donnee_enqueteur.banque_domiciliation),
                ('Libellé Guichet', donnee_enqueteur.libelle_guichet),
                ('Titulaire du Compte', donnee_enqueteur.titulaire_compte),
                ('Code Banque', donnee_enqueteur.code_banque),
                ('Code Guichet', donnee_enqueteur.code_guichet),
            ])
            
            # Section 14: Informations de décès
            if donnee_enqueteur.date_deces:
                add_table_section(doc, '14. Informations de Décès', [
                    ('Date de Décès', donnee_enqueteur.date_deces.strftime('%d/%m/%Y')),
                    ('N° Acte de Décès', donnee_enqueteur.numero_acte_deces),
                    ('Code INSEE Décès', donnee_enqueteur.code_insee_deces),
                    ('Code Postal Décès', donnee_enqueteur.code_postal_deces),
                    ('Localité Décès', donnee_enqueteur.localite_deces),
                ])
            
            # Section 15: Informations revenus
            add_table_section(doc, '15. Informations sur les Revenus', [
                ('Commentaires Revenus', donnee_enqueteur.commentaires_revenus),
                ('Montant Salaire', f"{donnee_enqueteur.montant_salaire:.2f} €" if donnee_enqueteur.montant_salaire else None),
                ('Période Versement Salaire', donnee_enqueteur.periode_versement_salaire),
                ('Fréquence Versement Salaire', donnee_enqueteur.frequence_versement_salaire),
            ])
            
            # Section 16: Autres revenus
            revenus_data = []
            if donnee_enqueteur.nature_revenu1:
                revenus_data.extend([
                    ('Nature Revenu 1', donnee_enqueteur.nature_revenu1),
                    ('Montant Revenu 1', f"{donnee_enqueteur.montant_revenu1:.2f} €" if donnee_enqueteur.montant_revenu1 else None),
                    ('Période Versement 1', donnee_enqueteur.periode_versement_revenu1),
                    ('Fréquence Versement 1', donnee_enqueteur.frequence_versement_revenu1),
                ])
            if donnee_enqueteur.nature_revenu2:
                revenus_data.extend([
                    ('Nature Revenu 2', donnee_enqueteur.nature_revenu2),
                    ('Montant Revenu 2', f"{donnee_enqueteur.montant_revenu2:.2f} €" if donnee_enqueteur.montant_revenu2 else None),
                    ('Période Versement 2', donnee_enqueteur.periode_versement_revenu2),
                    ('Fréquence Versement 2', donnee_enqueteur.frequence_versement_revenu2),
                ])
            if donnee_enqueteur.nature_revenu3:
                revenus_data.extend([
                    ('Nature Revenu 3', donnee_enqueteur.nature_revenu3),
                    ('Montant Revenu 3', f"{donnee_enqueteur.montant_revenu3:.2f} €" if donnee_enqueteur.montant_revenu3 else None),
                    ('Période Versement 3', donnee_enqueteur.periode_versement_revenu3),
                    ('Fréquence Versement 3', donnee_enqueteur.frequence_versement_revenu3),
                ])
            
            if revenus_data:
                add_table_section(doc, '16. Autres Revenus', revenus_data)
            
            # Section 17: Facturation
            add_table_section(doc, '17. Informations de Facturation', [
                ('N° Facture', donnee_enqueteur.numero_facture),
                ('Date Facture', donnee_enqueteur.date_facture.strftime('%d/%m/%Y') if donnee_enqueteur.date_facture else None),
                ('Montant Facture', f"{donnee_enqueteur.montant_facture:.2f} €" if donnee_enqueteur.montant_facture else None),
                ('Tarif Appliqué', f"{donnee_enqueteur.tarif_applique:.2f} €" if donnee_enqueteur.tarif_applique else None),
                ('Cumul Montants Précédents', f"{donnee_enqueteur.cumul_montants_precedents:.2f} €" if donnee_enqueteur.cumul_montants_precedents else None),
                ('Reprise Facturation', f"{donnee_enqueteur.reprise_facturation:.2f} €" if donnee_enqueteur.reprise_facturation else None),
                ('Remise Éventuelle', f"{donnee_enqueteur.remise_eventuelle:.2f} €" if donnee_enqueteur.remise_eventuelle else None),
            ])
            
            # Section 18: Mémos
            memos_data = []
            for i in range(1, 6):
                memo_attr = f'memo{i}'
                if hasattr(donnee_enqueteur, memo_attr):
                    memo_value = getattr(donnee_enqueteur, memo_attr)
                    if memo_value:
                        memos_data.append((f'Mémo {i}', memo_value))
            
            if memos_data:
                add_table_section(doc, '18. Mémos et Notes', memos_data)
        
        # Footer
        doc.add_paragraph()
        footer_para = doc.add_paragraph(f"Document généré le {datetime.datetime.now().strftime('%d/%m/%Y à %H:%M')}")
        footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in footer_para.runs:
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(128, 128, 128)
            run.font.italic = True
    
    return doc

@export_bp.route('/api/export/enquete/<int:enquete_id>', methods=['POST'])
def export_and_archive_enquete(enquete_id):
    """Exporte une enquête individuelle en Word et l'archive"""
    try:
        donnee = db.session.get(Donnee, enquete_id)
        if not donnee:
            return jsonify({"error": "Enquête non trouvée"}), 404
        
        if donnee.statut_validation != 'archive':
            return jsonify({"error": "Seules les enquêtes archivées peuvent être exportées"}), 400
        
        from models.enquete_archive import EnqueteArchive
        existing_archive = EnqueteArchive.query.filter_by(enquete_id=enquete_id).first()
        if existing_archive:
            return jsonify({"error": "Cette enquête a déjà été archivée"}), 400
        
        doc = generate_word_document([donnee])
        
        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        filename = f"Enquete_{donnee.numeroDossier}_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        
        archive = EnqueteArchive(
            enquete_id=enquete_id,
            nom_fichier=filename,
            utilisateur=request.json.get('utilisateur', 'Administrateur') if request.json else 'Administrateur'
        )
        db.session.add(archive)
        db.session.commit()
        
        logger.info(f"Enquête {enquete_id} exportée et archivée avec succès")
        
        return send_file(
            file_stream,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        db.session.rollback()
        logger.error(f"Erreur lors de l'export et archivage de l'enquête: {str(e)}")
        return jsonify({"error": f"Erreur lors de l'export: {str(e)}"}), 500

@export_bp.route('/api/archives', methods=['GET'])
def get_archives():
    """Récupère la liste des enquêtes archivées"""
    try:
        from models.enquete_archive import EnqueteArchive
        
        page = request.args.get('page', 1, type=int)
        per_page = request.args.get('per_page', 50, type=int)
        
        archives_query = db.session.query(
            EnqueteArchive, Donnee
        ).join(
            Donnee, EnqueteArchive.enquete_id == Donnee.id
        ).order_by(
            EnqueteArchive.date_export.desc()
        )
        
        archives_paginated = archives_query.paginate(page=page, per_page=per_page, error_out=False)
        
        result = []
        for archive, donnee in archives_paginated.items:
            result.append({
                'id': archive.id,
                'enquete_id': archive.enquete_id,
                'numeroDossier': donnee.numeroDossier,
                'nom': donnee.nom,
                'prenom': donnee.prenom,
                'nom_fichier': archive.nom_fichier,
                'date_export': archive.date_export.strftime('%Y-%m-%d %H:%M:%S'),
                'utilisateur': archive.utilisateur
            })
        
        return jsonify({
            'success': True,
            'data': result,
            'page': page,
            'per_page': per_page,
            'total': archives_paginated.total,
            'pages': archives_paginated.pages
        })
    except Exception as e:
        logger.error(f"Erreur lors de la récupération des archives: {str(e)}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@export_bp.route('/api/archives/<int:archive_id>', methods=['GET'])
def get_archive_file(archive_id):
    """Télécharge un fichier archivé"""
    try:
        from models.enquete_archive import EnqueteArchive
        
        archive = db.session.get(EnqueteArchive, archive_id)
        if not archive:
            return jsonify({"error": "Archive non trouvée"}), 404
        
        donnee = db.session.get(Donnee, archive.enquete_id)
        if not donnee:
            return jsonify({"error": "Enquête non trouvée"}), 404
        
        doc = generate_word_document([donnee])
        
        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        return send_file(
            file_stream,
            as_attachment=True,
            download_name=archive.nom_fichier if archive.nom_fichier else f"enquete_{archive_id}.docx",
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        logger.error(f"Erreur lors de la récupération du fichier archivé: {str(e)}")
        return jsonify({"error": f"Erreur lors de la récupération: {str(e)}"}), 500

@export_bp.route('/api/exports/validated', methods=['GET'])
def get_enquetes_validees_pour_export():
    """
    Récupère toutes les enquêtes validées (statut='validee') prêtes pour l'export EOS uniquement
    Ces enquêtes ont été validées par l'admin mais pas encore archivées
    FILTRE: Exclut les enquêtes PARTNER (client_id du client PARTNER)
    """
    try:
        page = request.args.get('page', 1, type=int)
        per_page = request.args.get('per_page', 100, type=int)
        
        # Récupérer l'ID du client PARTNER pour l'exclure
        from models.client import Client
        partner_client = Client.query.filter_by(code='PARTNER').first()
        partner_id = partner_client.id if partner_client else None
        
        # Récupérer les enquêtes avec statut_validation = 'validee'
        # EXCLURE les enquêtes PARTNER
        enquetes_query = db.session.query(
            Donnee, DonneeEnqueteur
        ).join(
            DonneeEnqueteur, Donnee.id == DonneeEnqueteur.donnee_id
        ).filter(
            Donnee.statut_validation == 'validee',
            DonneeEnqueteur.code_resultat.in_(['P', 'H', 'N', 'Z', 'I', 'Y'])
        )
        
        # Exclure PARTNER si le client existe
        if partner_id:
            enquetes_query = enquetes_query.filter(Donnee.client_id != partner_id)
        
        enquetes_query = enquetes_query.order_by(Donnee.updated_at.desc())
        
        # Pagination
        enquetes_paginated = enquetes_query.paginate(page=page, per_page=per_page, error_out=False)
        
        result = []
        for donnee, donnee_enqueteur in enquetes_paginated.items:
            enqueteur = db.session.get(Enqueteur, donnee.enqueteurId) if donnee.enqueteurId else None
            enqueteur_nom = f"{enqueteur.nom} {enqueteur.prenom}" if enqueteur else "Non assigné"
            
            result.append({
                'id': donnee.id,
                'numeroDossier': donnee.numeroDossier,
                'referenceDossier': donnee.referenceDossier,
                'nom': donnee.nom,
                'prenom': donnee.prenom,
                'typeDemande': donnee.typeDemande,
                'enqueteurId': donnee.enqueteurId,
                'enqueteurNom': enqueteur_nom,
                'code_resultat': donnee_enqueteur.code_resultat,
                'elements_retrouves': donnee_enqueteur.elements_retrouves,
                'updated_at': donnee.updated_at.strftime('%Y-%m-%d %H:%M:%S') if donnee.updated_at else None,
                'statut_validation': donnee.statut_validation
            })
        
        return jsonify({
            'success': True,
            'data': result,
            'page': page,
            'per_page': per_page,
            'total': enquetes_paginated.total,
            'pages': enquetes_paginated.pages
        })
    except Exception as e:
        logger.error(f"Erreur lors de la récupération des enquêtes validées: {str(e)}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


def generate_eos_export_line(donnee, donnee_enqueteur, enqueteur, export_date, facturation=None):
    """
    Génère une ligne d'export au format "Réponses EOS" (fixed-width 2618 chars + CRLF).

    STRICTEMENT CONFORME AU CAHIER DES CHARGES + FICHIER EXEMPLE:
    - Longueur fixe: EXACTEMENT 2618 caractères (hors CRLF) + \\r\\n
    - Champs identifiants: valeurs EXACTES transmises par EOS (pas d'IDs internes)
    - État civil: priorité corrections (e.*_corrige), sinon valeurs d (d.*)
    - Facturation: calculée depuis tarifs_eos (date export + code tarif)
    - Date retour: date_retour réelle, fallback aujourd'hui
    - REVENUS (274 chars): commentaires(128) + salaire(14) + revenus1-3(44 chacun)
    - MÉMOS (1256 chars): memo1-4(64) + memo5(1000)
    - numeroCompte (11) et RIB (2): TOUJOURS VIDES (espaces)
    - Montants facturation (8 chars): POINT décimal, padding espaces
    - Montants revenus (10 chars): POINT décimal, padding espaces
    - Entiers facturation (8 chars): padding espaces
    - Format Windows CRLF (\\r\\n)
    - Validation longueur STRICTE (len=2618, sinon reject)

    Args:
        donnee: Objet Donnee (table donnees)
        donnee_enqueteur: Objet DonneeEnqueteur (table donnees_enqueteur)
        enqueteur: Objet Enqueteur (non utilisé, compatibilité signature)
        export_date: Date d'export (utilisée pour facturation)
        facturation: Objet EnqueteFacturation (optionnel, non utilisé actuellement)

    Returns:
        str: Ligne formatée + CRLF, ou None si erreur/champs manquants
    """
    # === VALIDATION DES CHAMPS OBLIGATOIRES ===
    # Note: Pour compatibilité avec données existantes, on utilise des fallbacks minimaux
    # mais on logue les problèmes au lieu de bloquer tout l'export
    champs_obligatoires = {
        # Identifiants (depuis donnee)
        'numeroDossier': donnee.numeroDossier,
        'referenceDossier': donnee.referenceDossier,
        'numeroInterlocuteur': donnee.numeroInterlocuteur,
        'guidInterlocuteur': donnee.guidInterlocuteur,
        'typeDemande': donnee.typeDemande,
        'numeroDemande': donnee.numeroDemande,
        'forfaitDemande': donnee.forfaitDemande,
        # Résultat (depuis donnee_enqueteur)
        'code_resultat': donnee_enqueteur.code_resultat if hasattr(donnee_enqueteur, 'code_resultat') else None,
        'elements_retrouves': donnee_enqueteur.elements_retrouves if hasattr(donnee_enqueteur, 'elements_retrouves') else None,
    }

    champs_manquants = [nom for nom, valeur in champs_obligatoires.items() if not valeur]
    if champs_manquants:
        logger.warning(f"Enquête ID={donnee.id} ignorée - champs obligatoires manquants: {', '.join(champs_manquants)}")
        return None  # Ignorer cette ligne au lieu de bloquer tout l'export

    # Déterminer si c'est une contestation
    est_contestation = (donnee.typeDemande == 'CON')

    # === CONSTRUCTION DE LA LIGNE SELON LE CAHIER DES CHARGES ===
    fields = []

    # 1. N° DOSSIER (10) - EXACTEMENT celui transmis par EOS
    fields.append(format_alphanum_eos(donnee.numeroDossier, 10))

    # 2. RÉFÉRENCE DOSSIER (15) - EXACTEMENT celle transmise par EOS
    fields.append(format_alphanum_eos(donnee.referenceDossier, 15))

    # 3. NUMÉRO INTERLOCUTEUR (12) - EXACTEMENT celui transmis par EOS
    fields.append(format_alphanum_eos(donnee.numeroInterlocuteur, 12))

    # 4. GUID INTERLOCUTEUR (36) - EXACTEMENT celui transmis par EOS
    fields.append(format_alphanum_eos(donnee.guidInterlocuteur, 36))

    # 5. TYPE DE DEMANDE (3) - ENQ ou CON selon la donnée
    fields.append(format_alphanum_eos(donnee.typeDemande, 3))

    # 6. NUMÉRO DE DEMANDE (11) - EXACTEMENT celui transmis par EOS
    fields.append(format_alphanum_eos(donnee.numeroDemande, 11))

    # 7-8. NUMÉRO DEMANDE CONTESTÉE/INITIALE (11 + 11)
    # Rempli uniquement pour les contestations (CON)
    if est_contestation:
        fields.append(format_alphanum_eos(donnee.numeroDemandeContestee or '', 11))
        fields.append(format_alphanum_eos(donnee.numeroDemandeInitiale or '', 11))
    else:
        fields.append(format_alphanum_eos('', 11))
        fields.append(format_alphanum_eos('', 11))

    # 9. FORFAIT DE DEMANDE (16) - EXACTEMENT celui transmis par EOS
    fields.append(format_alphanum_eos(donnee.forfaitDemande, 16))

    # 10. DATE DE RETOUR ESPÉRÉ (10) - EXACTEMENT celle transmise par EOS
    fields.append(format_date_eos(donnee.dateRetourEspere))

    # 11. QUALITÉ / CIVILITÉ (10) - PRIORITÉ CORRECTION
    qualite = donnee_enqueteur.qualite_corrigee if hasattr(donnee_enqueteur, 'qualite_corrigee') and donnee_enqueteur.qualite_corrigee else donnee.qualite
    fields.append(format_alphanum_eos(qualite or '', 10))

    # 12-13. NOM, PRÉNOM (30 + 20) - PRIORITÉ CORRECTION
    nom = donnee_enqueteur.nom_corrige if hasattr(donnee_enqueteur, 'nom_corrige') and donnee_enqueteur.nom_corrige else donnee.nom
    prenom = donnee_enqueteur.prenom_corrige if hasattr(donnee_enqueteur, 'prenom_corrige') and donnee_enqueteur.prenom_corrige else donnee.prenom
    fields.append(format_alphanum_eos(nom or '', 30))
    fields.append(format_alphanum_eos(prenom or '', 20))

    # 14-17. DATE/LIEU/CP/PAYS NAISSANCE (10 + 50 + 10 + 32)
    fields.append(format_date_eos(donnee.dateNaissance))
    fields.append(format_alphanum_eos(donnee.lieuNaissance or '', 50))

    # CP/Pays naissance - PRIORITÉ CORRECTION
    cp_naissance = donnee_enqueteur.code_postal_naissance_corrige if hasattr(donnee_enqueteur, 'code_postal_naissance_corrige') and donnee_enqueteur.code_postal_naissance_corrige else donnee.codePostalNaissance
    pays_naissance = donnee_enqueteur.pays_naissance_corrige if hasattr(donnee_enqueteur, 'pays_naissance_corrige') and donnee_enqueteur.pays_naissance_corrige else donnee.paysNaissance
    fields.append(format_alphanum_eos(cp_naissance or '', 10))
    fields.append(format_alphanum_eos(pays_naissance or '', 32))

    # 18. NOM PATRONYMIQUE (30) - PRIORITÉ CORRECTION
    nom_patronymique = donnee_enqueteur.nom_patronymique_corrige if hasattr(donnee_enqueteur, 'nom_patronymique_corrige') and donnee_enqueteur.nom_patronymique_corrige else donnee.nomPatronymique
    fields.append(format_alphanum_eos(nom_patronymique or '', 30))

    # 19. DATE DE RETOUR (10) - Date retour enquêteur (fallback aujourd'hui)
    date_retour = donnee_enqueteur.date_retour if hasattr(donnee_enqueteur, 'date_retour') and donnee_enqueteur.date_retour else datetime.date.today()
    fields.append(format_date_eos(date_retour))

    # 20-21. CODE RÉSULTAT, ÉLÉMENTS RETROUVÉS (1 + 10)
    fields.append(format_alphanum_eos(donnee_enqueteur.code_resultat or '', 1))
    fields.append(format_alphanum_eos(donnee_enqueteur.elements_retrouves or '', 10))

    # 22. FLAG ÉTAT CIVIL ERRONÉ (1)
    fields.append(format_alphanum_eos(donnee_enqueteur.flag_etat_civil_errone or '', 1))

    # 23-29. FACTURATION (9 + 10 + 8*5)
    # N° facture: toujours vide (9 espaces)
    fields.append(format_alphanum_eos('', 9))
    
    # Date facture: date dexport (10 chars DD/MM/YYYY)
    fields.append(format_date_eos(export_date))

    # Récupérer le tarif depuis tarifs_eos
    tarif_montant = get_tarif_from_eos(donnee, donnee_enqueteur, export_date)
    
    # Montant facturé: tarif lookup (8 chars avec point et espaces)
    fields.append(format_montant_8(tarif_montant))
    
    # Tarif appliqué: même valeur que montant facturé (8 chars)
    fields.append(format_montant_8(tarif_montant))

    # Cumul montants précédents: toujours 0.00 (8 chars)
    fields.append(format_montant_8(0.0))

    # Reprise facturation: toujours 0 (8 chars entier avec espaces)
    fields.append(format_int_8(0))
    
    # Remise éventuelle: toujours 0 (8 chars entier avec espaces)
    fields.append(format_int_8(0))


    # 30-34. DÉCÈS (10 + 10 + 5 + 10 + 32)
    fields.append(format_date_eos(donnee_enqueteur.date_deces))
    fields.append(format_alphanum_eos(donnee_enqueteur.numero_acte_deces or '', 10))
    fields.append(format_alphanum_eos(donnee_enqueteur.code_insee_deces or '', 5))
    fields.append(format_alphanum_eos(donnee_enqueteur.code_postal_deces or '', 10))
    fields.append(format_alphanum_eos(donnee_enqueteur.localite_deces or '', 32))

    # 35-38. ADRESSE (32*4 + 10 + 32 + 32)
    fields.append(format_alphanum_eos(donnee_enqueteur.adresse1 or '', 32))
    fields.append(format_alphanum_eos(donnee_enqueteur.adresse2 or '', 32))
    fields.append(format_alphanum_eos(donnee_enqueteur.adresse3 or '', 32))
    fields.append(format_alphanum_eos(donnee_enqueteur.adresse4 or '', 32))
    fields.append(format_alphanum_eos(donnee_enqueteur.code_postal or '', 10))
    fields.append(format_alphanum_eos(donnee_enqueteur.ville or '', 32))
    fields.append(format_alphanum_eos(donnee_enqueteur.pays_residence or '', 32))

    # 39-40. TÉLÉPHONES (15 + 15)
    fields.append(format_alphanum_eos(donnee_enqueteur.telephone_personnel or '', 15))
    fields.append(format_alphanum_eos(donnee_enqueteur.telephone_chez_employeur or '', 15))

    # 41-47. EMPLOYEUR (32 + 15 + 15 + 32*4 + 10 + 32 + 32)
    fields.append(format_alphanum_eos(donnee_enqueteur.nom_employeur or '', 32))
    fields.append(format_alphanum_eos(donnee_enqueteur.telephone_employeur or '', 15))
    fields.append(format_alphanum_eos(donnee_enqueteur.telecopie_employeur or '', 15))
    fields.append(format_alphanum_eos(donnee_enqueteur.adresse1_employeur or '', 32))
    fields.append(format_alphanum_eos(donnee_enqueteur.adresse2_employeur or '', 32))
    fields.append(format_alphanum_eos(donnee_enqueteur.adresse3_employeur or '', 32))
    fields.append(format_alphanum_eos(donnee_enqueteur.adresse4_employeur or '', 32))
    fields.append(format_alphanum_eos(donnee_enqueteur.code_postal_employeur or '', 10))
    fields.append(format_alphanum_eos(donnee_enqueteur.ville_employeur or '', 32))
    fields.append(format_alphanum_eos(donnee_enqueteur.pays_employeur or '', 32))

    # 48-54. BANQUE (32 + 30 + 32 + 5 + 5 + 11 + 2)
    fields.append(format_alphanum_eos(donnee_enqueteur.banque_domiciliation or '', 32))
    fields.append(format_alphanum_eos(donnee_enqueteur.libelle_guichet or '', 30))
    fields.append(format_alphanum_eos(donnee_enqueteur.titulaire_compte or '', 32))
    fields.append(format_alphanum_eos(donnee_enqueteur.code_banque or '', 5))
    fields.append(format_alphanum_eos(donnee_enqueteur.code_guichet or '', 5))
    fields.append(format_alphanum_eos('', 11))  # Numéro compte - TOUJOURS VIDE
    fields.append(format_alphanum_eos('', 2))   # RIB - TOUJOURS VIDE

    # === BLOC REVENUS (274 chars: 128+14+44*3) - FORMAT "RÉPONSES" ===
    # 55. Commentaires revenus (128)
    fields.append(format_alphanum_eos(donnee_enqueteur.commentaires_revenus if hasattr(donnee_enqueteur, 'commentaires_revenus') else '', 128))

    # 56-58. Salaire (10+2+2 = 14)
    fields.append(format_montant_10(donnee_enqueteur.montant_salaire if hasattr(donnee_enqueteur, 'montant_salaire') else None))
    fields.append(format_numeric_eos(donnee_enqueteur.periode_versement_salaire if hasattr(donnee_enqueteur, 'periode_versement_salaire') else None, 2))
    fields.append(format_alphanum_eos(donnee_enqueteur.frequence_versement_salaire if hasattr(donnee_enqueteur, 'frequence_versement_salaire') else '', 2))

    # 59-62. Revenu 1 (30+10+2+2 = 44)
    fields.append(format_alphanum_eos(donnee_enqueteur.nature_revenu1 if hasattr(donnee_enqueteur, 'nature_revenu1') else '', 30))
    fields.append(format_montant_10(donnee_enqueteur.montant_revenu1 if hasattr(donnee_enqueteur, 'montant_revenu1') else None))
    fields.append(format_numeric_eos(donnee_enqueteur.periode_versement_revenu1 if hasattr(donnee_enqueteur, 'periode_versement_revenu1') else None, 2))
    fields.append(format_alphanum_eos(donnee_enqueteur.frequence_versement_revenu1 if hasattr(donnee_enqueteur, 'frequence_versement_revenu1') else '', 2))

    # 63-66. Revenu 2 (30+10+2+2 = 44)
    fields.append(format_alphanum_eos(donnee_enqueteur.nature_revenu2 if hasattr(donnee_enqueteur, 'nature_revenu2') else '', 30))
    fields.append(format_montant_10(donnee_enqueteur.montant_revenu2 if hasattr(donnee_enqueteur, 'montant_revenu2') else None))
    fields.append(format_numeric_eos(donnee_enqueteur.periode_versement_revenu2 if hasattr(donnee_enqueteur, 'periode_versement_revenu2') else None, 2))
    fields.append(format_alphanum_eos(donnee_enqueteur.frequence_versement_revenu2 if hasattr(donnee_enqueteur, 'frequence_versement_revenu2') else '', 2))

    # 67-70. Revenu 3 (30+10+2+2 = 44)
    fields.append(format_alphanum_eos(donnee_enqueteur.nature_revenu3 if hasattr(donnee_enqueteur, 'nature_revenu3') else '', 30))
    fields.append(format_montant_10(donnee_enqueteur.montant_revenu3 if hasattr(donnee_enqueteur, 'montant_revenu3') else None))
    fields.append(format_numeric_eos(donnee_enqueteur.periode_versement_revenu3 if hasattr(donnee_enqueteur, 'periode_versement_revenu3') else None, 2))
    fields.append(format_alphanum_eos(donnee_enqueteur.frequence_versement_revenu3 if hasattr(donnee_enqueteur, 'frequence_versement_revenu3') else '', 2))

    # === BLOC MÉMOS (1128 chars) - FORMAT "RÉPONSES" ===
    # 72-76. Mémos (64 + 64 + 64 + 64 + 1000)
    fields.append(format_alphanum_eos(donnee_enqueteur.memo1 if hasattr(donnee_enqueteur, 'memo1') else '', 64))
    fields.append(format_alphanum_eos(donnee_enqueteur.memo2 if hasattr(donnee_enqueteur, 'memo2') else '', 64))
    fields.append(format_alphanum_eos(donnee_enqueteur.memo3 if hasattr(donnee_enqueteur, 'memo3') else '', 64))
    fields.append(format_alphanum_eos(donnee_enqueteur.memo4 if hasattr(donnee_enqueteur, 'memo4') else '', 64))
    fields.append(format_alphanum_eos(donnee_enqueteur.memo5 if hasattr(donnee_enqueteur, 'memo5') else '', 1000))

    # Joindre tous les champs
    line = ''.join(fields)

    # VALIDATION HARD DE LA LONGUEUR - STRICTEMENT 2618 caractères
    # Calcul: Identif(135) + EtatCivil(192) + Résultat(22) + Facturation(59) +
    #         Décès(67) + Adresse(202) + Tél(30) + Employeur(264) + Banque(117) +
    #         Revenus(274) + Mémos(1256) = 2618 caractères
    EXPECTED_LENGTH = 2618

    if len(line) != EXPECTED_LENGTH:
        logger.error(f"ERREUR LONGUEUR export EOS: enquête ID={donnee.id}")
        logger.error(f"  Attendu: {EXPECTED_LENGTH} chars")
        logger.error(f"  Obtenu:  {len(line)} chars")
        logger.error(f"  Différence: {len(line) - EXPECTED_LENGTH:+d} chars")
        return None

    # Ajouter CRLF Windows - IMPORTANT: newline='' dans open() pour préserver
    return line + '\r\n'


@export_bp.route('/api/exports/create-batch', methods=['POST'])
def create_export_batch():
    """
    Crée un export groupé de toutes les enquêtes validées au format EOS
    - Génère un fichier texte (.txt) à longueur fixe selon le cahier des charges EOS
    - Sauvegarde le fichier sur disque
    - Marque les enquêtes comme 'archivee'
    - Crée une entrée ExportBatch pour tracker l'export
    - Retourne le fichier pour téléchargement
    """
    try:
        # Récupérer les paramètres
        data = request.json or {}
        utilisateur = data.get('utilisateur', 'Administrateur')
        
        # Limite maximale pour un seul export EOS (performance)
        MAX_EXPORT_EOS_LIMIT = 5000
        
        # Compter le nombre total d'enquêtes validées
        total_count = Donnee.query.filter_by(statut_validation='validee').count()
        
        if total_count == 0:
            return jsonify({
                'success': False,
                'error': 'Aucune enquête validée à exporter'
            }), 404
        
        # Avertir si dépassement de la limite
        if total_count > MAX_EXPORT_EOS_LIMIT:
            logger.warning(f"Export EOS limité : {total_count} enquêtes disponibles, mais limite fixée à {MAX_EXPORT_EOS_LIMIT}")
        
        # Récupérer les enquêtes validées (avec limite)
        donnees = Donnee.query.filter_by(statut_validation='validee')\
                              .order_by(Donnee.created_at.asc())\
                              .limit(MAX_EXPORT_EOS_LIMIT).all()
        
        if not donnees:
            return jsonify({
                'success': False,
                'error': 'Aucune enquête validée à exporter'
            }), 404
        
        logger.info(f"Création d'un export EOS de {len(donnees)} enquêtes sur {total_count} validée(s) par {utilisateur}")
        
        # Créer le dossier exports/batches s'il n'existe pas
        base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        batches_dir = os.path.join(base_dir, 'exports', 'batches')
        os.makedirs(batches_dir, exist_ok=True)
        
        # Créer un nom de fichier selon le format EOS : XXXExp_AAAAMMJJ.txt
        export_date = datetime.date.today()  # Date d'export pour facturation
        date_str = export_date.strftime('%Y%m%d')
        filename = f"{CODE_PRESTATAIRE}Exp_{date_str}.txt"
        filepath_full = os.path.join(batches_dir, filename)
        
        # Générer le contenu du fichier au format EOS
        lines = []
        exported_ids = []  # IDs des enquêtes RÉELLEMENT exportées
        skipped_count = 0

        for donnee in donnees:
            # Récupérer les données enquêteur
            donnee_enqueteur = DonneeEnqueteur.query.filter_by(donnee_id=donnee.id).first()
            if not donnee_enqueteur:
                logger.warning(f"Pas de données enquêteur pour l'enquête {donnee.id}, ignorée")
                skipped_count += 1
                continue

            # Récupérer l'enquêteur (pour compatibilité signature)
            enqueteur = None
            if donnee.enqueteurId:
                enqueteur = db.session.get(Enqueteur, donnee.enqueteurId)

            # TODO: Récupérer facturation si disponible (pour fallback montants)
            # from models.enquete_facturation import EnqueteFacturation
            # facturation = EnqueteFacturation.query.filter_by(donnee_enqueteur_id=donnee_enqueteur.id).first()
            facturation = None  # Temporaire, en attendant modèle

            # Générer la ligne au format EOS "Réponses" avec export_date
            line = generate_eos_export_line(donnee, donnee_enqueteur, enqueteur, export_date, facturation)

            # Ignorer les lignes avec champs obligatoires manquants OU longueur invalide
            if line is None:
                skipped_count += 1
                continue

            # Ligne valide → ajouter à l'export
            lines.append(line)
            exported_ids.append(donnee.id)  # Tracer uniquement les enquêtes exportées
        
        if not lines:
            return jsonify({
                'success': False,
                'error': 'Aucune enquête avec données enquêteur à exporter'
            }), 404
        
        # Écrire le fichier avec encodage Windows (CP1252) et fin de ligne CR+LF
        # Utiliser errors='replace' pour gérer les caractères non supportés
        with open(filepath_full, 'w', encoding='cp1252', newline='', errors='replace') as f:
            f.writelines(lines)
        
        # Calculer la taille du fichier
        file_size = os.path.getsize(filepath_full)
        
        # Chemin relatif depuis le dossier exports/
        filepath_relative = os.path.join('batches', filename)
        
        # Récupérer le client_id (toutes les enquêtes exportées ensemble ont le même client)
        client_id = donnees[0].client_id if donnees else None

        # Créer l'entrée ExportBatch (UNIQUEMENT avec les IDs réellement exportés)
        export_batch = ExportBatch(
            client_id=client_id,
            filename=filename,
            filepath=filepath_relative,
            file_size=file_size,
            enquete_count=len(exported_ids),  # Nombre réel exporté
            utilisateur=utilisateur
        )
        export_batch.set_enquete_ids_list(exported_ids)  # UNIQUEMENT les IDs exportés
        db.session.add(export_batch)

        # Marquer UNIQUEMENT les enquêtes RÉELLEMENT EXPORTÉES comme archivées
        for donnee_id in exported_ids:
            donnee = db.session.get(Donnee, donnee_id)
            if donnee:
                donnee.statut_validation = 'archivee'
                donnee.add_to_history(
                    'archivage',
                    f'Enquête exportée au format EOS Réponses dans {filename} par {utilisateur}',
                    utilisateur
                )
        
        db.session.commit()

        # Log du résultat avec info sur les lignes ignorées
        if skipped_count > 0:
            logger.warning(f"Export EOS créé avec {skipped_count} enquête(s) ignorée(s) (champs obligatoires manquants): {filename} ({len(lines)} lignes exportées, {file_size} octets)")
        else:
            logger.info(f"Export EOS créé avec succès: {filename} ({len(lines)} lignes, {file_size} octets)")
        
        # Retourner le fichier pour téléchargement
        return send_file(
            filepath_full,
            as_attachment=True,
            download_name=filename,
            mimetype='text/plain; charset=cp1252'
        )
        
    except Exception as e:
        db.session.rollback()
        logger.error(f"Erreur lors de la création de l'export EOS: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({
            'success': False,
            'error': f"Erreur lors de la création de l'export: {str(e)}"
        }), 500


@export_bp.route('/api/exports/batches', methods=['GET'])
def get_export_batches():
    """
    Récupère la liste des exports batch créés
    Paramètres optionnels:
    - client_id: filtre par client (MULTI-CLIENT)
    - page: numéro de page (défaut: 1)
    - per_page: nombre d'éléments par page (défaut: 50)
    """
    try:
        page = request.args.get('page', 1, type=int)
        per_page = request.args.get('per_page', 50, type=int)
        client_id = request.args.get('client_id', None, type=int)
        
        # Construire la requête avec filtre optionnel par client
        batches_query = ExportBatch.query
        
        # MULTI-CLIENT: Filtrer par client_id si spécifié
        if client_id is not None:
            batches_query = batches_query.filter_by(client_id=client_id)
        
        batches_query = batches_query.order_by(ExportBatch.created_at.desc())
        batches_paginated = batches_query.paginate(page=page, per_page=per_page, error_out=False)
        
        result = [batch.to_dict() for batch in batches_paginated.items]
        
        return jsonify({
            'success': True,
            'data': result,
            'page': page,
            'per_page': per_page,
            'total': batches_paginated.total,
            'pages': batches_paginated.pages
        })
    except Exception as e:
        logger.error(f"Erreur lors de la récupération des batches: {str(e)}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@export_bp.route('/api/exports/batches/<int:batch_id>/download', methods=['GET'])
def download_export_batch(batch_id):
    """
    Télécharge un fichier d'export batch
    """
    try:
        batch = db.session.get(ExportBatch, batch_id)
        if not batch:
            return jsonify({
                'success': False,
                'error': 'Export batch non trouvé'
            }), 404
        
        # Construire le chemin complet du fichier
        base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        filepath_full = os.path.join(base_dir, 'exports', batch.filepath)
        
        # Vérifier que le fichier existe
        if not os.path.exists(filepath_full):
            logger.error(f"Fichier d'export batch non trouvé: {filepath_full}")
            return jsonify({
                'success': False,
                'error': 'Fichier d\'export non trouvé sur le disque'
            }), 404
        
        # Envoyer le fichier
        return send_file(
            filepath_full,
            as_attachment=True,
            download_name=batch.filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        logger.error(f"Erreur lors du téléchargement du batch {batch_id}: {str(e)}")
        return jsonify({
            'success': False,
            'error': f"Erreur lors du téléchargement: {str(e)}"
        }), 500


def register_export_routes(app):
    """
    Enregistre les routes d'export dans l'application
    """
    app.register_blueprint(export_bp)