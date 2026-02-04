"""
Service d'export pour le client PARTNER
Génère les fichiers Word (.docx) et Excel (.xls) selon le cahier des charges PARTNER
"""
import os
import logging
from datetime import datetime
from io import BytesIO
import tempfile

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import xlwt
    XLS_AVAILABLE = True
except ImportError:
    XLS_AVAILABLE = False

from extensions import db
from models.models import Donnee
from models.models_enqueteur import DonneeEnqueteur
from models.export_batch import ExportBatch
from models.tarifs import TarifClient
from models.partner_models import PartnerCaseRequest
from services.partner_tarif_resolver import PartnerTarifResolver
from utils import NanCleaner

logger = logging.getLogger(__name__)


class PartnerExportService:
    """Service pour gérer les exports PARTNER (Word + Excel)"""
    
    def __init__(self, client_id):
        self.client_id = client_id
        
    def _get_report_number(self, export_type):
        """
        Génère le numéro de rapport incrémental pour le jour
        export_type: 'enquete_positive', 'contestation_positive', etc.
        """
        today = datetime.now().date()
        
        # Compter les exports du même type créés aujourd'hui
        count = ExportBatch.query.filter(
            ExportBatch.client_id == self.client_id,
            db.func.date(ExportBatch.created_at) == today,
            ExportBatch.filename.like(f'%{export_type}%')
        ).count()
        
        return count + 1
    
    def _get_batch_total(self, donnee):
        """
        Retourne le nombre total de dossiers du même batch d'import
        """
        if donnee.fichier_id:
            batch_total = Donnee.query.filter_by(
                fichier_id=donnee.fichier_id,
                client_id=self.client_id
            ).count()
            return batch_total
        return 1
    
    def _get_montant_from_tarif(self, tarif_lettre):
        """
        Récupère le montant du tarif PARTNER depuis la lettre
        Retourne 0 si le tarif n'est pas trouvé
        """
        if not tarif_lettre:
            logger.warning(f"Tarif lettre vide pour le client {self.client_id}")
            return 0
        
        # Normaliser la lettre (trim + uppercase)
        code_lettre = str(tarif_lettre).strip().upper()
        
        # Chercher le tarif dans TarifClient
        tarif = TarifClient.query.filter_by(
            client_id=self.client_id,
            code_lettre=code_lettre,
            actif=True
        ).first()
        
        if tarif:
            logger.debug(f"Tarif trouvé pour lettre '{code_lettre}': {float(tarif.montant)}€")
            return float(tarif.montant)
        else:
            logger.warning(f"Tarif PARTNER non trouvé pour lettre '{code_lettre}' (client_id={self.client_id})")
            return 0
    
    def _format_reference_enquete(self, donnee, batch_total):
        """
        Formate la référence pour une enquête: {DATE_ENVOI dd.MM}/{BATCH_TOTAL} {TARIF}
        """
        date_envoi = donnee.datedenvoie or donnee.created_at.date()
        date_str = date_envoi.strftime('%d.%m') if date_envoi else '00.00'
        tarif = donnee.tarif_lettre or ''
        return f"{date_str}/{batch_total} {tarif}"
    
    def _format_reference_contestation(self, donnee, batch_total):
        """
        Formate la référence pour une contestation: {DATE_JOUR dd.MM}/{BATCH_TOTAL}
        """
        date_jour = donnee.date_jour or donnee.date_contestation or donnee.created_at.date()
        date_str = date_jour.strftime('%d.%m') if date_jour else '00.00'
        return f"{date_str}/{batch_total}"
    
    def _normalize_address(self, address_dict):
        """
        Normalise une adresse pour comparaison
        """
        parts = []
        if address_dict.get('adresse1'):
            parts.append(address_dict['adresse1'].upper().strip())
        if address_dict.get('adresse2'):
            parts.append(address_dict['adresse2'].upper().strip())
        if address_dict.get('adresse3'):
            parts.append(address_dict['adresse3'].upper().strip())
        if address_dict.get('adresse4'):
            parts.append(address_dict['adresse4'].upper().strip())
        if address_dict.get('code_postal'):
            parts.append(str(address_dict['code_postal']).strip())
        if address_dict.get('ville'):
            parts.append(address_dict['ville'].upper().strip())
        
        return ' '.join(parts)
    
    def _is_address_confirmed(self, donnee, donnee_enqueteur):
        """
        Détermine si l'adresse résultat est une confirmation ou une nouvelle adresse
        """
        if not donnee_enqueteur:
            return False
        
        # Adresse importée
        addr_import = {
            'adresse1': donnee.adresse1,
            'adresse2': donnee.adresse2,
            'adresse3': donnee.adresse3,
            'adresse4': donnee.adresse4,
            'code_postal': donnee.codePostal,
            'ville': donnee.ville
        }
        
        # Adresse résultat
        addr_result = {
            'adresse1': donnee_enqueteur.adresse1,
            'adresse2': donnee_enqueteur.adresse2,
            'adresse3': donnee_enqueteur.adresse3,
            'adresse4': donnee_enqueteur.adresse4,
            'code_postal': donnee_enqueteur.code_postal,
            'ville': donnee_enqueteur.ville
        }
        
        return self._normalize_address(addr_import) == self._normalize_address(addr_result)
    
    def generate_enquetes_positives_word(self, enquetes):
        """
        Génère le document Word pour les enquêtes positives PARTNER
        Format simple: paragraphes uniquement, titre+référence répétés par dossier
        """
        if not DOCX_AVAILABLE:
            raise Exception("python-docx n'est pas disponible")

        doc = Document()

        section = doc.sections[0]
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

        report_no = self._get_report_number('enquete_positive')
        date_export = datetime.now().strftime('%d/%m/%Y')

        for idx, donnee_obj in enumerate(enquetes):
            if idx > 0:
                doc.add_paragraph()
                doc.add_paragraph()

            donnee = NanCleaner(donnee_obj)
            donnee_enqueteur = NanCleaner(donnee_obj.donnee_enqueteur) if donnee_obj.donnee_enqueteur else None
            batch_total = self._get_batch_total(donnee)

            # Titre (répété par dossier)
            p_titre = doc.add_paragraph(f"Rapport positif du {date_export} no {report_no}")
            p_titre.runs[0].bold = True

            # Référence du batch
            reference = self._format_reference_enquete(donnee, batch_total)
            doc.add_paragraph(reference)

            # Ligne identité: NOM  PRENOM (+ NJF) alignée avec NO numeroDossier
            nom = (donnee.nom or '').strip()
            prenom = (donnee.prenom or '').strip()
            identity = f"{nom}  {prenom}"
            if donnee.nomPatronymique:
                identity += f"  NJF {donnee.nomPatronymique}"
            num = str(donnee.numeroDossier or '0')
            doc.add_paragraph(identity.ljust(35) + f"NO {num}")

            # Naissance
            if donnee.dateNaissance:
                jour = donnee.dateNaissance.day
                mois = donnee.dateNaissance.month
                annee = donnee.dateNaissance.year
                lieu = donnee.lieuNaissance or ''
                doc.add_paragraph(f"Ne le {jour}/{mois}/{annee} a {lieu}".rstrip())

            # Ligne vide avant bloc résultat
            doc.add_paragraph()

            if donnee_enqueteur:
                has_address = any([
                    donnee_enqueteur.adresse1,
                    donnee_enqueteur.adresse2,
                    donnee_enqueteur.adresse3,
                    donnee_enqueteur.adresse4
                ])
                has_employer = bool(donnee_enqueteur.nom_employeur)

                # Bloc adresse résultat
                if has_address:
                    is_confirmed = self._is_address_confirmed(donnee, donnee_enqueteur)
                    doc.add_paragraph("CONFIRMATION ADRESSE:" if is_confirmed else "NOUVELLE ADRESSE:")

                    if donnee_enqueteur.adresse1:
                        doc.add_paragraph(donnee_enqueteur.adresse1)
                    if donnee_enqueteur.adresse2:
                        doc.add_paragraph(donnee_enqueteur.adresse2)
                    if donnee_enqueteur.adresse3:
                        doc.add_paragraph(donnee_enqueteur.adresse3)
                    if donnee_enqueteur.adresse4:
                        doc.add_paragraph(donnee_enqueteur.adresse4)

                    cp = (donnee_enqueteur.code_postal or '').strip()
                    ville = (donnee_enqueteur.ville or '').strip()
                    if cp or ville:
                        doc.add_paragraph(f"{cp} {ville}".strip())

                    tel = donnee_enqueteur.telephone_personnel or donnee_enqueteur.telephone_chez_employeur
                    if tel:
                        doc.add_paragraph(f"Tel: {tel}")

                # Bloc employeur
                if has_employer:
                    if has_address:
                        doc.add_paragraph()
                    doc.add_paragraph("EMPLOYEUR:")
                    doc.add_paragraph(donnee_enqueteur.nom_employeur)

                    if donnee_enqueteur.adresse1_employeur:
                        doc.add_paragraph(donnee_enqueteur.adresse1_employeur)
                    if donnee_enqueteur.adresse2_employeur:
                        doc.add_paragraph(donnee_enqueteur.adresse2_employeur)
                    if donnee_enqueteur.adresse3_employeur:
                        doc.add_paragraph(donnee_enqueteur.adresse3_employeur)
                    if donnee_enqueteur.adresse4_employeur:
                        doc.add_paragraph(donnee_enqueteur.adresse4_employeur)

                    if donnee_enqueteur.telephone_employeur:
                        doc.add_paragraph(f"tel {donnee_enqueteur.telephone_employeur}")

                    if donnee_enqueteur.memo3:
                        doc.add_paragraph(donnee_enqueteur.memo3)

        output = BytesIO()
        doc.save(output)
        output.seek(0)

        return output
    
    def generate_enquetes_positives_excel(self, enquetes):
        """
        Génère le fichier Excel (.xls) pour les enquêtes positives
        Colonnes exactes selon le cahier des charges
        """
        if not XLS_AVAILABLE:
            raise Exception("xlwt n'est pas disponible")
        
        # Colonnes exactes
        columns = [
            "NUM", "DATE BUTOIR", "DATE ENVOI", "TARIF", "NOM", "PRENOM", "NJF",
            "JOUR", "MOIS", "ANNEE NAISSANCE", "LIEUNAISSANCE", "PAYSNAISSANCE",
            "NOM2", "PRENOM2", "JOUR2", "MOIS2", "ANNEE NAISSANCE2", "LIEUNAISSANCE2", "PAYSNAISSANCE2",
            "ADRESSE", "CP", "VILLE", "PAYS", "TEL", "TEL2", "AUTRE'", "AUTRE2",
            "TITULAIRE", "CODEBANQUE", "COMPTE",
            "EMPLOYEUR", "ADRESSE_EMPLOYEUR",
            "INSTRUCTIONS", "RECHERCHE",
            "Proximite", "Date naissance (MAJ)", "Lieu naissance (MAJ)", 
            "Adresse 1", "Adresse 2", "Adresse 3", "Adresse 4",
            "Code postal", "Ville", "Pays",
            "Telephone 1", "Telephone 2", "Portable 1", "Portable 2",
            "Montant facture", "memo",
            "Nom banque", "Code Banque", "Code guichet",
            "Adresse 1 banque", "Adresse 2 banque", "Adresse 3 banque", "Adresse 4 banque",
            "Telepone banque", "memo banque",
            "Nom employeur", "Adresse 1 employeur", "Adresse 2 employeur",
            "Adresse 3 employeur", "Adresse 4 employeur",
            "Telephone employeur", "Memo employeur",
            "Resultat"
        ]
        
        workbook = xlwt.Workbook()
        sheet = workbook.add_sheet('Enquetes Positives')
        
        # Style pour l'en-tête
        header_style = xlwt.XFStyle()
        font = xlwt.Font()
        font.bold = True
        header_style.font = font
        
        # Écrire les en-têtes
        for col_idx, col_name in enumerate(columns):
            sheet.write(0, col_idx, col_name, header_style)
        
        # Écrire les données
        for row_idx, donnee_obj in enumerate(enquetes, start=1):
            donnee = NanCleaner(donnee_obj)
            donnee_enqueteur = NanCleaner(donnee_obj.donnee_enqueteur) if donnee_obj.donnee_enqueteur else None
            
            # Construire la ligne de données
            row_data = []
            
            # Colonnes import
            row_data.append(donnee.numeroDossier or '')
            row_data.append(donnee.date_butoir.strftime('%d/%m/%Y') if donnee.date_butoir else '')
            row_data.append(donnee.datedenvoie.strftime('%d/%m/%Y') if donnee.datedenvoie else '')
            row_data.append(donnee.tarif_lettre or '')
            row_data.append(donnee.nom or '')
            row_data.append(donnee.prenom or '')
            row_data.append(donnee.nomPatronymique or '')  # NJF
            
            # Date naissance (JOUR, MOIS, ANNEE NAISSANCE)
            if donnee.dateNaissance:
                row_data.append(donnee.dateNaissance.day)
                row_data.append(donnee.dateNaissance.month)
                row_data.append(donnee.dateNaissance.year)
            else:
                row_data.extend(['', '', ''])
            
            row_data.append(donnee.lieuNaissance or '')
            row_data.append(donnee.paysNaissance or '')
            
            # Champs NOM2, PRENOM2, etc. (vides pour l'instant)
            row_data.extend(['', '', '', '', '', '', ''])
            
            # Adresse import
            row_data.append(donnee.adresse1 or '')
            row_data.append(donnee.codePostal or '')
            row_data.append(donnee.ville or '')
            row_data.append(donnee.paysResidence or '')
            row_data.append(donnee.telephonePersonnel or '')
            row_data.append(donnee.telephoneEmployeur or '')
            row_data.extend(['', ''])  # AUTRE', AUTRE2
            
            # Banque
            row_data.append(donnee.titulaireCompte or '')
            row_data.append(donnee.codeBanque or '')
            row_data.append(donnee.numeroCompte or '')
            
            # Employeur import
            row_data.append(donnee.nomEmployeur or '')
            row_data.append('')  # ADRESSE_EMPLOYEUR (peut être combinée)
            
            # INSTRUCTIONS et RECHERCHE
            row_data.append(donnee.instructions or '')
            row_data.append(donnee.recherche or '')
            
            # Résultats enquêteur
            if donnee_enqueteur:
                # Proximite : utiliser elements_retrouves ("Confirmation par qui" dans l'UI)
                row_data.append(donnee_enqueteur.elements_retrouves or '')  # Proximite
                
                # Date et lieu de naissance mis à jour (depuis Donnee)
                if donnee.dateNaissance_maj:
                    row_data.append(donnee.dateNaissance_maj.strftime('%d/%m/%Y'))
                else:
                    row_data.append('')
                row_data.append(donnee.lieuNaissance_maj or '')
                
                row_data.append(donnee_enqueteur.adresse1 or '')
                row_data.append(donnee_enqueteur.adresse2 or '')
                row_data.append(donnee_enqueteur.adresse3 or '')
                row_data.append(donnee_enqueteur.adresse4 or '')
                row_data.append(donnee_enqueteur.code_postal or '')
                row_data.append(donnee_enqueteur.ville or '')
                row_data.append(donnee_enqueteur.pays_residence or '')
                row_data.append(donnee_enqueteur.telephone_personnel or '')
                row_data.append(donnee_enqueteur.telephone_chez_employeur or '')
                row_data.append('')  # Portable 1 (n'existe pas)
                row_data.append('')  # Portable 2 (n'existe pas)
                # Calculer le montant avec le tarif combiné (lettre + demandes)
                try:
                    resolver = PartnerTarifResolver()
                    montant = resolver.resolve_tarif(donnee, self.client_id)
                    if montant is None:
                        # Fallback : utiliser le tarif simple si pas de règle combinée
                        montant = self._get_montant_from_tarif(donnee.tarif_lettre)
                        logger.warning(f"Pas de tarif combiné pour dossier {donnee.id}, utilisation tarif simple: {montant}€")
                except Exception as e:
                    logger.error(f"Erreur calcul tarif combiné pour dossier {donnee.id}: {e}")
                    montant = self._get_montant_from_tarif(donnee.tarif_lettre)
                row_data.append(montant)
                row_data.append(donnee_enqueteur.memo1 or '')
                
                # Banque enquêteur
                row_data.append(donnee_enqueteur.banque_domiciliation or '')
                row_data.append(donnee_enqueteur.code_banque or '')
                row_data.append(donnee_enqueteur.code_guichet or '')
                row_data.append('')  # Adresse 1 banque (n'existe pas)
                row_data.append('')  # Adresse 2 banque (n'existe pas)
                row_data.append('')  # Adresse 3 banque (n'existe pas)
                row_data.append('')  # Adresse 4 banque (n'existe pas)
                row_data.append('')  # Telephone banque (n'existe pas)
                row_data.append('')  # Memo banque (n'existe pas)
                
                # Employeur enquêteur
                row_data.append(donnee_enqueteur.nom_employeur or '')
                row_data.append(donnee_enqueteur.adresse1_employeur or '')
                row_data.append(donnee_enqueteur.adresse2_employeur or '')
                row_data.append(donnee_enqueteur.adresse3_employeur or '')
                row_data.append(donnee_enqueteur.adresse4_employeur or '')
                row_data.append(donnee_enqueteur.telephone_employeur or '')
                row_data.append(donnee_enqueteur.memo3 or '')  # Utiliser memo3 pour memo employeur
            else:
                # Remplir avec des valeurs vides
                row_data.extend([''] * 27)
            
            # Resultat
            row_data.append('POS')
            
            # Écrire la ligne
            for col_idx, value in enumerate(row_data):
                sheet.write(row_idx, col_idx, value)
        
        # Sauvegarder dans un BytesIO
        output = BytesIO()
        workbook.save(output)
        output.seek(0)
        
        return output
    
    def generate_enquetes_negatives_excel(self, enquetes):
        """
        Génère le fichier Excel (.xls) pour les enquêtes négatives
        Colonnes: nom, prenom, reference, dossier, memo
        Génère un fichier avec headers même si enquetes est vide (robustesse)
        """
        if not XLS_AVAILABLE:
            raise Exception("xlwt n'est pas disponible")
        
        columns = ["nom", "prenom", "reference", "dossier", "memo"]
        
        workbook = xlwt.Workbook()
        sheet = workbook.add_sheet('Enquetes Negatives')
        
        # Style pour l'en-tête
        header_style = xlwt.XFStyle()
        font = xlwt.Font()
        font.bold = True
        header_style.font = font
        
        # Écrire les en-têtes (toujours, même si 0 enquêtes)
        for col_idx, col_name in enumerate(columns):
            sheet.write(0, col_idx, col_name, header_style)
        
        # Écrire les données (si présentes)
        logger.info(f"Génération Excel enquêtes négatives: {len(enquetes)} lignes")
        for row_idx, donnee_obj in enumerate(enquetes, start=1):
            donnee = NanCleaner(donnee_obj)
            batch_total = self._get_batch_total(donnee)
            reference = self._format_reference_enquete(donnee, batch_total)
            
            row_data = [
                donnee.nom or '',
                donnee.prenom or '',
                reference,
                donnee.numeroDossier or '',
                ''  # memo vide
            ]
            
            for col_idx, value in enumerate(row_data):
                sheet.write(row_idx, col_idx, value)
        
        # Sauvegarder dans un BytesIO
        output = BytesIO()
        workbook.save(output)
        output.seek(0)
        
        return output
    
    def generate_contestations_positives_word(self, contestations):
        """
        Génère le document Word pour les contestations positives
        Format similaire aux enquêtes positives mais avec ajustements
        """
        if not DOCX_AVAILABLE:
            raise Exception("python-docx n'est pas disponible")
        
        doc = Document()
        
        # Configuration des marges
        section = doc.sections[0]
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        
        report_no = self._get_report_number('contestation_positive')
        date_export = datetime.now().strftime('%d/%m/%Y')
        
        for idx, donnee_obj in enumerate(contestations):
            if idx > 0:
                # Ajouter 2 lignes vides entre les blocs
                doc.add_paragraph()
                doc.add_paragraph()

            donnee = NanCleaner(donnee_obj)
            donnee_enqueteur = NanCleaner(donnee_obj.donnee_enqueteur) if donnee_obj.donnee_enqueteur else None
            batch_total = self._get_batch_total(donnee)
            
            # Titre
            titre = f"Rapport contestation positif du {date_export} no {report_no}"
            p_titre = doc.add_paragraph(titre)
            p_titre.runs[0].bold = True
            
            # Référence
            reference = self._format_reference_contestation(donnee, batch_total)
            doc.add_paragraph(reference)
            
            # URGENT si applicable
            if donnee.urgence == '1' or donnee.urgence == 'O':
                p_urgent = doc.add_paragraph("URGENT")
                p_urgent.runs[0].bold = True
            
            # Identité (nom_complet ou nom + prenom)
            num = donnee.numeroDossier or '0'
            nom_complet = donnee.nom_complet or f"{donnee.nom or ''} {donnee.prenom or ''}".strip()
            identite = f"{nom_complet} NO {num}"
            doc.add_paragraph(identite)
            
            # Naissance si dispo
            if donnee.dateNaissance:
                jour = donnee.dateNaissance.day
                mois = donnee.dateNaissance.month
                annee = donnee.dateNaissance.year
                lieu = donnee.lieuNaissance or ''
                naissance = f"Ne le {jour}/{mois}/{annee} a {lieu}"
                doc.add_paragraph(naissance)
            
            # Adresse résultat
            if donnee_enqueteur and (donnee_enqueteur.adresse1 or donnee_enqueteur.adresse2 or
                                    donnee_enqueteur.adresse3 or donnee_enqueteur.adresse4):
                doc.add_paragraph("NOUVELLE ADRESSE:")
                if donnee_enqueteur.adresse1:
                    doc.add_paragraph(donnee_enqueteur.adresse1)
                if donnee_enqueteur.adresse2:
                    doc.add_paragraph(donnee_enqueteur.adresse2)
                if donnee_enqueteur.adresse3:
                    doc.add_paragraph(donnee_enqueteur.adresse3)
                if donnee_enqueteur.adresse4:
                    doc.add_paragraph(donnee_enqueteur.adresse4)
                
                # Code postal + Ville
                if donnee_enqueteur.code_postal or donnee_enqueteur.ville:
                    cp = donnee_enqueteur.code_postal or ''
                    ville = donnee_enqueteur.ville or ''
                    doc.add_paragraph(f"{cp} {ville}".strip())
                
                # Téléphone
                if donnee_enqueteur.telephone_personnel:
                    doc.add_paragraph(f"Tel: {donnee_enqueteur.telephone_personnel}")
                elif donnee_enqueteur.telephone_chez_employeur:
                    doc.add_paragraph(f"Tel: {donnee_enqueteur.telephone_chez_employeur}")
                
                # Mémo
                if donnee_enqueteur.memo1:
                    doc.add_paragraph(donnee_enqueteur.memo1)
                if donnee_enqueteur.memo2:
                    doc.add_paragraph(donnee_enqueteur.memo2)
        
        # Sauvegarder dans un BytesIO
        output = BytesIO()
        doc.save(output)
        output.seek(0)
        
        return output
    
    def generate_contestations_negatives_excel(self, contestations):
        """
        Génère le fichier Excel (.xls) pour les contestations négatives
        Colonnes: nom, prenom, reference, dossier, memo
        Génère un fichier avec headers même si contestations est vide (robustesse)
        """
        if not XLS_AVAILABLE:
            raise Exception("xlwt n'est pas disponible")
        
        columns = ["nom", "prenom", "reference", "dossier", "memo"]
        
        workbook = xlwt.Workbook()
        sheet = workbook.add_sheet('Contestations Negatives')
        
        # Style pour l'en-tête
        header_style = xlwt.XFStyle()
        font = xlwt.Font()
        font.bold = True
        header_style.font = font
        
        # Écrire les en-têtes (toujours, même si 0 contestations)
        for col_idx, col_name in enumerate(columns):
            sheet.write(0, col_idx, col_name, header_style)
        
        # Écrire les données (si présentes)
        logger.info(f"Génération Excel contestations négatives: {len(contestations)} lignes")
        for row_idx, donnee_obj in enumerate(contestations, start=1):
            donnee = NanCleaner(donnee_obj)
            batch_total = self._get_batch_total(donnee)
            reference = self._format_reference_contestation(donnee, batch_total)
            
            # nom = nom_complet ou nom standard
            nom = donnee.nom_complet or donnee.nom or ''
            
            # prenom = "TRES URGENT" si urgent, sinon vide
            prenom = "TRES URGENT" if (donnee.urgence == '1' or donnee.urgence == 'O') else ''
            
            row_data = [
                nom,
                prenom,
                reference,
                donnee.numeroDossier or '',
                'NEGATIF'
            ]
            
            for col_idx, value in enumerate(row_data):
                sheet.write(row_idx, col_idx, value)
        
        # Sauvegarder dans un BytesIO
        output = BytesIO()
        workbook.save(output)
        output.seek(0)
        
        return output
    
    def create_export_batch(self, enquete_ids, export_type, filename, filepath, file_size):
        """
        Crée un enregistrement ExportBatch et marque les enquêtes comme exportées et archivées
        """
        try:
            # Créer le batch
            batch = ExportBatch(
                client_id=self.client_id,
                filename=filename,
                filepath=filepath,
                file_size=file_size,
                enquete_count=len(enquete_ids),
                enquete_ids=','.join(str(id) for id in enquete_ids),
                created_at=datetime.now()
            )
            
            db.session.add(batch)
            db.session.flush()  # Pour obtenir l'ID du batch
            
            # Marquer les enquêtes comme exportées ET archivées
            now = datetime.now()
            for enquete_id in enquete_ids:
                donnee = Donnee.query.get(enquete_id)
                if donnee:
                    donnee.exported = True
                    donnee.exported_at = now
                    donnee.statut_validation = 'archivee'
            
            db.session.commit()
            
            logger.info(f"Batch d'export créé: {len(enquete_ids)} enquêtes archivées")
            
            return batch
            
        except Exception as e:
            db.session.rollback()
            logger.error(f"Erreur lors de la création du batch d'export: {str(e)}")
            raise

